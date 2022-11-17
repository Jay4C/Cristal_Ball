import unittest
from docx import Document


class CV_with_Docx(unittest.TestCase):
    # ok
    def test_1(self):
        print('test_1')

        document = Document()

        document.add_heading('Document Title', 0)

        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='Intense Quote')

        document.add_paragraph(
            'first item in unordered list', style='List Bullet'
        )
        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        document.add_page_break()

        document.save('CV_Word\\demo.docx')

    # ok
    def test_cv_developpeur_python(self):
        print("test_cv_developpeur_python")

        document = Document()

        document.add_heading('Développeur Python', 0)

        # Compétences
        document.add_heading('Compétences', level=1)

        # Front-end
        document.add_paragraph('Front-end', style='Intense Quote')

        records_front_end = [
            'Django web framework / Langage de balisage HTML / Langage CSS',
            'Librairie Bootstrap / LibrairieLeaflet.js / Librairie Vis.js / Librairie Artyom.js / '
            'LibrairieFlowchart.js / Librairie Babylon.js / Librairie Chart.js / Librairie MathJax',
            'Serveur Apache/Serveur Nginx'
        ]

        for record in records_front_end:
            document.add_paragraph(record)
        # Front-end

        # Back-end
        document.add_paragraph('Back-end', style='Intense Quote')

        records_back_end = [
            'Langage de programmation Python / Langage C++ avec la carte Arduino / '
            'Framework Flask / LibrairieSQLAlchemy / Librairie unittests / '
            'Librairie BeautifulSoup / LibrairieSelenium WebDriver / Librairie pymysql.cursors / Librairierequests / '
            'Librairie pdfkit / Librairiexlswriter / Librairiexlrd / Librairie pywinauto / LibrairieSkidl / Librairie '
            'PyPDF2 / Librairiereportlab / Librairie pylatex / Librairiesympy / Librairie numpy / Librairie pandas / '
            'Librairie matplotlib / Librairieahkab / Librairieemail / Librairiesmtplib / Librairieschemdraw / '
            'Librairie paramiko / Librairieimaplib / Librairieimap_tools / Librairie minimalmodbus / '
            'Librairie pyModbusTCP',
            'Serveur de bases de données relationnelles MySQL /Serveur de bases de données documents MongoDB / '
            'Serveur de bases de données graphe de Neo4j',
            'Exploration de données / Automatisation des processus robotisés',
            'Carte embarquée Raspberry Pi / Système d’exploitation Raspbian',
            'API (Programmableweb) / Logiciel de gestion deversion Git / GitHub Docker',
            'Logiciel FreeCAD / Logiciel KiCAD / Logiciel SonarQube',
            'Windows / Linux / Ubuntu / Bash / Serveur privé virtuel'
        ]

        for record in records_back_end:
            document.add_paragraph(record)
        # Back-end

        # Management
        document.add_paragraph('Management', style='Intense Quote')

        records_management = [
            "Brainstorming / Logiciel FreeMind / Logiciel Pencil / Diagrammes UML / Logiciel ArgoUML",
            "Diagramme de Gantt / Logiciel GanttProject / Planification des étapes d'un projet",
            "Management de la connaissance/ Management de projet / Management des opérations / "
            "Management du changement / Management de la qualité",
            "Analyse des besoins du client / Établissement d'un cahier des charges",
            "Rédaction d'un support technique/ Rédaction d'une base de données deconnaissances",
            'Droit français : "www.legifrance.gouv.fr" / Droiteuropéen : '
            '"https://europa.eu/european-union/law/find-legislation_fr" / Droit international : "www.lexadin.nl"',
            'Anglais écritet oral [TOEIC : 840 points ; Date d’expiration : 11 Juillet 2020]'
        ]

        for record in records_management:
            document.add_paragraph(record)
        # Management

        # Compétences

        # Expériences
        document.add_heading('Expériences', level=1)

        experiences = [
            {
                'intitule': 'Développeur Python',
                'periode': 'Mai 2019 - En cours',
                'missions': [
                    "Développement d'un site internet (Python / Django web framework)",
                    "Développement d'applications Python pour créer des pièces mécaniques de précision pour les "
                    "technologies de l'hydrogène avec lelogiciel FreeCAD (Python)",
                    "Développement d'applications Python pour l'automatisation de processus robotisés (Python / "
                    "Selenium)",
                    "Développement d'une application Python pour récupérer toutes les parcelles convenables pour "
                    "injecter du méthane de synthèse dans un réseau de gaz naturel en France (Python / API Open data)",
                    "Développement de tests unitaires pour programmer des robots d'exploration de données "
                    "(Python / BeautifulSoup)",
                    "Développement de tests unitaires pour faire du mailing (Python / smtplib)",
                    "Développement de tests unitaires pour préparer des requêtes vers des API (Python / API)",
                    "Développement de tests unitaires pour faire des diagrammes de circuits électroniques "
                    "(Python / Librairie schemdraw)",
                    "Développement de tests unitaires pour fabriquer des circuits électroniques "
                    "(Python / Librairie Skidl)",
                    "Développement detests unitaires pour simuler des circuits électroniques (Python / Librairie "
                    "ahkab / Librairie pyspice)"
                ]
            },
            {
                'intitule': 'Développeur Python',
                'periode': 'Octobre 2016 - Avril 2019 [30 mois]',
                'missions': [
                    "Développement de tests unitaires pour paramétrer des trajectoires de bras-robot avec la carte "
                    "Raspberry Pi (Python)",
                    "Développement de tests unitaires pour concevoir des circuits électroniques sous le logiciel KiCAD "
                    "(Python)",
                    "Développement de tests unitaires pour converir des séries d'images en vidéos (Python)"
                ]
            },
            {
                'intitule': 'Développeur Python',
                'periode': 'Avril 2015 - Juillet 2015 [4 mois] et Février 2016 - Juillet 2016 [6 mois]',
                'missions': [
                    "Développement de programmes informatiques pour collecter des données énergétiques sur des "
                    "chaudières avec la carte Raspberry Pi grâce au protocole Modbus (Python - Librairie minimalmodbus)",
                    "Développement de programmes informatiques pour piloter des chaudières à distanceavec la carte "
                    "Raspberry Pi (Python - Librairie pyModbusTCP)",
                    "Développement de programmes informatiques pour stocker des données énergétiques issues de "
                    "chaudières dans un serveur de base de données MySQL (Python)"
                ]
            }
        ]

        for experience in experiences:
            document.add_paragraph(
                "Intitulé : " + experience.get('intitule') + " / Période : " + experience.get('periode'),
                style='Intense Quote'
            )

            for mission in experience.get('missions'):
                document.add_paragraph(mission)
        # Expériences

        # Formations
        document.add_heading('Formations', level=1)

        formations = [
            {
                'intitule': "Diplôme d'ingénieur généraliste",
                'periode': 'Septembre 2013 - Août 2016 [3 ans]',
                'centre': "École Supérieure d'Ingénieurs Léonard de Vinci – Paris La Défense",
                'enseignements': [
                    "Nouvelles énergies / Informatique/ Finance/ Mécanique numérique",
                    "Veille technologique / Physique/ Mathématiques / Economie / Management / Marketing / Comptabilité"
                ]
            },
            {
                'intitule': ' Classes préparatoires aux grandes écoles PTSI-PT',
                'periode': ' Août 2011 - Août 2013 [2 ans]',
                'centre': "Lycée Lislet Geofrroy à l'Ile de La Réunion",
                'enseignements': [
                    "Mathématiques / Chimie / Physique / Technologie / Sciences de l'ingénieur"
                ]
            },
            {
                'intitule': ' Baccalauréat scientifique',
                'periode': ' Août 2010 - Août 2011 [1 an]',
                'centre': "Lycée Roland Garros à l'Ile de La Réunion",
                'enseignements': [
                    "Sciences de l'ingénieur / Mathématiques / Physique / Chimie / Anglais / Français / Histoire"
                ]
            },
        ]

        for formation in formations:
            document.add_paragraph(
                "Intitulé : " + formation.get('intitule') +
                " / Période : " + formation.get('periode') +
                " / Centre de formation : " + formation.get('centre'),
                style='Intense Quote'
            )

            for enseignement in formation.get('enseignements'):
                document.add_paragraph(enseignement)
        # Formations

        document.save('CV_Word\\CV_J_A_[Developpeur_Python].docx')


if __name__ == '__main__':
    unittest.main()
