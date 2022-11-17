import unittest
import pdfkit


class CV_with_HTML(unittest.TestCase):
    # ok
    def test_cv(self):
        print("test_cv")

        # Reporting
        body = """
<!doctype html>
<html lang="en">

<head>
    <!-- Required meta tags -->
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

    <!-- Bootstrap CSS -->
    <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
        integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

    <title>CV de Monsieur  - Développeur Full-Stack</title>
</head>

<body>
    <br>
    <div class="container">
        <div class="card">
            <div class="card-header text-center">
                <h2>
                    Développeur Full-Stack
                </h2>
            </div>
            <div class="card-body text-center">

                <!-- Informations personnelles -->
                <table class="table table-striped table-bordered">
                    <thead>
                        <tr>
                            <th scope="col">Identité</th>
                            <th scope="col">Adresse postale</th>
                            <th scope="col">Téléphone</th>
                            <th scope="col">E-mail</th>
                            <th scope="col">&Acirc;ge</th>
                            <th scope="col">Nationalit&eacute;</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr>
                            <td>
                                Monsieur Jason
                                <br>
                                ALOYAU
                            </td>
                            <td>
                                31 Avenue de Ségur
                                <br>
                                75007 Paris
                            </td>
                            <td></td>
                            <td>@holomorphe.com</td>
                            <td>27 ans</td>
                            <td>Française</td>
                        </tr>
                    </tbody>
                </table>

                <!--  Compétences -->
                <div class="card">
                    <div class="card-header">
                        <h4>
                            Comp&eacute;tences
                        </h4>
                    </div>
                    <div class="card-body">
                        <!-- Anglais -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Anglais
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Anglais écrit et oral [TOEIC : 840 points ; Date d’expiration : 11 Juillet 2020]
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Anglais -->

                        <br>

                        <!-- Langages de programmation de prédilection -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Langages de programmation de prédilection
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>Langage Python / Langage Java / Langage JavaScript</li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Langages de programmation de prédilection -->

                        <br>

                        <!-- Bureautique -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Bureautique
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>Logiciels Microsoft Excel, Word, PowerPoint, Access, Outlook</li>
                                            <li>Logiciels Open office Calc, Writer</li>
                                            <li>
                                                Logiciel IntelliJ pour Java /
                                                Logiciel PyCharm pour Python
                                            </li>
                                            <li>
                                                Logiciel Pencil pour faire des schémas simplifiés / 
                                                Logiciel GanttProject pour le management de projet
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Bureautique -->

                        <br>

                        <!-- Front-end -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Front-end
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Langage de balisage HTML / 
                                                Langage CSS /
                                                Framework JQuery / 
                                                Framework Django / 
                                                Framework MJML pour le responsive email
                                            </li>
                                            <li>
                                                Logiciel Visual Studio Code /
                                                Framework Angular 7 /
                                                Librairie Bootstrap / 
                                                Librairie Leaflet.js
                                            </li>
                                            <li>
                                                Librairie Vis.js / 
                                                Librairie Artyom.js /
                                                Librairie Flowchart.js
                                            </li>
                                            <li>
                                                Serveur Tomcat / 
                                                Serveur Apache /
                                                Serveur Nginx
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Front-end -->

                        <br>

                        <!-- Back-end -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Back-end
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Langage Python /
                                                Framework Flask /
                                                Librairie BeautifulSoup /
                                                Librairie pdfkit /
                                                Librairie xlswriter /
                                                Librairie xlrd /
                                                Librairie pywinauto /
                                                Librairie Hyperledger Sawtooth
                                            </li>
                                            <li>
                                                Langage Java /
                                                Gestionnaire de dépendances Maven / 
                                                Gestionnaire de dépendances Gradle /
                                                Librairie Spring /
                                                Librairie Jsoup /
                                                Librairie Selenium WebDriver
                                            </li>
                                            <li>
                                                Langage JavaScript /
                                                Langage Visual Basic /
                                                Langage C# /
                                                Langage TypeScript
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Back-end -->

                        <br>

                        <!-- Tests unitaires et qualité -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Tests unitaires et qualité
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Package Unittests avec Python / 
                                                Package JUnit avec Java / 
                                                Mockito avec Java / 
                                                WireMock avec Java
                                            </li>
                                            <li>
                                                Logiciel SonarQube
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Tests unitaires et qualité -->

                        <br>

                        <!-- Systèmes de bases de données -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Systèmes de bases de données
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Serveur de bases de données MySQL / 
                                                Logiciel MySQL Workbench
                                            </li>
                                            <li>
                                                Serveur de bases de données graphe de Neo4j / 
                                                Serveur de bases de données documents MongoDB
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Systèmes de bases de données -->

                        <br>

                        <!-- Outils de déploiement -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Outils de déploiement
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Docker / Git
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Outils de déploiement -->

                        <br>

                        <!-- Intelligence artificielle -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Intelligence artificielle
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Exploration de données / 
                                                Gestion de la connaissance / 
                                                Automatisation de processus robotisés /
                                                RPA
                                            </li>
                                            <li>
                                                Cloud computing [Open data; Autodesk Forge; Netheos] / 
                                                Reconnaissance vocale / 
                                                Vision par ordinateur
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Intelligence artificielle -->

                        <br>

                        <!-- Mobile -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Mobile
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Système d'exploitation Android / 
                                                Logiciel Android Studio
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Mobile -->

                        <br>

                        <!-- Systèmes embarqués -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Systèmes embarqués
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Carte embarquée Raspberry Pi / 
                                                Système d’exploitation Raspbian
                                            </li>
                                            <li>
                                                Cartes embarquées Arduino Uno et Arduino Yun /
                                                Système d’exploitation OpenWRT
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Systèmes embarqués -->

                        <br>

                        <!-- Conception assistée par ordinateur -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Conception assistée par ordinateur
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Logiciel FreeCAD / 
                                                Logiciel Skidl
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Conception assistée par ordinateur -->

                        <br>

                        <!-- Gestion de version -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Gestion de version
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Logiciel de gestion de version Git / 
                                                Gestion de mon compte GitHub : https://github.com/Jay4C
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Gestion de version -->

                        <br>

                        <!-- Réseaux et Systèmes -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Réseaux et systèmes
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Système d’exploitation Windows / 
                                                Système d’exploitation Linux
                                            </li>
                                            <li>
                                                Protocoles informatiques [HTTP, HTTPS, TCP, IPv4, IPv6, SSH, FTP, 
                                                SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, Modbus RTU]
                                            </li>
                                            <li>
                                                Librairie MinimalModbus avec Python / 
                                                Librairie Jlibmodbus avec Java
                                            </li>
                                            <li>
                                                Routeur 4G / 
                                                Port Forwarding / 
                                                DynDNS / 
                                                Dongle 3G
                                            </li>
                                            <li>
                                                Logiciel Wireshark / 
                                                Logiciel PuTTY / 
                                                Logiciel WinSCP /
                                                Logiciel Postman /
                                                Logiciel TeamViewer / 
                                                Logiciel VirtualBox /
                                                Logiciel Tor /
                                                Logiciel Angry IP Scanner
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Réseaux et Systèmes -->

                        <br>

                        <!-- Gestion de projet et communications -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Gestion de projet et communications
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Brainstorming / 
                                                Logiciel FreeMind / 
                                                Logiciel Pencil / 
                                                Diagrammes UML / 
                                                Logiciel ArgoUML
                                            </li>
                                            <li>
                                                Logiciel GanttProject / 
                                                Planification des étapes d'un projet
                                            </li>
                                            <li>
                                                Management de projet / 
                                                Management des opérations / 
                                                Management Agile / 
                                                Management de la qualité
                                            </li>
                                            <li>
                                                Analyse des besoins du client / 
                                                Établissement d’un cahier des charges
                                            </li>
                                            <li>
                                                Rédaction d’un support technique / 
                                                Rédaction d'une base de données de connaissances
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Gestion de projet et communications -->

                        <br>

                        <!-- Droit -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Droit
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Code général des impôts / 
                                                Code monétaire et financier / 
                                                Code de l'environnement / 
                                                Code de l'énergie
                                            </li>
                                            <li>
                                                Code des relations entre le public et l’administration / 
                                                Code de la construction et de l'habitation
                                            </li>
                                            <li>
                                                Code de l'urbanisme /
                                                Code des postes et des communications électroniques /                                                 
                                            </li>
                                            <li>
                                                Site internet "www.legifrance.gouv.fr" /
                                                Site internet "www.lexadin.nl" /
                                                Site internet "https://europa.eu/european-union/law/find-legislation_fr"
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Droit -->
                    </div>
                </div>
                <!-- Compétences -->

                <br>

                <!-- Expériences -->
                <div class="card">
                    <div class="card-header">
                        <h4>
                            Expériences
                        </h4>
                    </div>
                    <div class="card-body">
                        <!-- Holomorphe -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Développement informatique polyvalent
                                        <br>
                                        Période : Mai 2020 – En cours
                                        <br>
                                        Statut professionnel : Président
                                        <br>
                                        Société : Holomorphe - Paris
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement d'applications Python pour créer des pièces mécaniques 
                                                de précision pour les technologies de l'hydrogène avec le logiciel
                                                FreeCAD disponibles sur https://github.com/Jay4C/Python-Macros-For-FreeCAD
                                            </li>
                                            <li>
                                                Développement d'applications Python pour l'automatisation de processus 
                                                robotisés disponibles sur https://github.com/Jay4C/Web-Automation
                                            </li>
                                            <li>
                                                Développement d'une application Python pour récupérer toutes 
                                                les parcelles convenables pour injecter du méthane de synthèse dans un 
                                                réseau de distribution de gaz naturel en France
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Holomorphe -->

                        <br>

                        <!-- Société PASS Technologie -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Développeur informatique polyvalent
                                        <br>
                                        Période : Février 2019 - Avril 2019 [3 mois]
                                        <br>
                                        Statut professionnel : Salarié
                                        <br>
                                        Société : PASS Technologie - Paris
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement d'une application pour le BIM (Building Information
                                                Modeling)
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Société Pass Tech -->

                        <br>

                        <!-- Kriir -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Ingénieur d'études et développement
                                        <br>
                                        Période : Septembre – Novembre 2018 [3 mois]
                                        <br>
                                        Statut professionnel : Salarié
                                        <br>
                                        Société : Kriir - Ermont
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement d’une application web en Java pour utiliser la signature
                                                électronique
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Kriir -->

                        <br>

                        <!-- Refuel -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Lecture et écriture sur un automate en modbus
                                        <br>
                                        Période : Juin - Juillet 2018 [2 mois]
                                        <br>
                                        Statut professionnel : Indépendant
                                        <br>
                                        Société : Refuel S.A.S. - Trappes
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement de plusieurs fonctionnalités permettant de lire et
                                                d’écrire sur un automate programmable via le
                                                protocole Modbus TCP avec le langage Java en liaison WiFi
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Refuel  -->

                        <br>

                        <!-- SETEC SMART EFFICIENCY -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Génération automatique de rapport énergétique
                                        <br>
                                        Période : Novembre - Décembre 2017 [2 mois]
                                        <br>
                                        Statut professionnel : Indépendant
                                        <br>
                                        Société : Setec Smart Efficiency - Paris
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développer un programme Java pour générer un rapport énergétique
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- SETEC SMART EFFICIENCY -->

                        <br>

                        <!--  SOCIÉTÉ GAMELOFT -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Vérification de la pertinence des données
                                        <br>
                                        Période : Juin - Juillet 2017 [2 mois]
                                        <br>
                                        Statut professionnel : Indépendant
                                        <br>
                                        Société : Gameloft - Paris
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement d’une macro Visual Basic pour un fichier Excel 
                                                permettant de contrôler la cohérence des données saisies dans 
                                                plusieurs fichiers Excel et consolider ces données dans un seul 
                                                fichier Excel pour les ressources humaines
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- SOCIÉTÉ GAMELOFT -->

                        <br>
                    </div>
                </div>
                <!-- Expériences -->

                <br>

                <!-- Formations -->
                <div class="card">
                    <div class="card-header">
                        <h4>
                            Formations
                        </h4>
                    </div>
                    <div class="card-body">
                        <!-- Formation Java -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Formation Java / JEE
                                        <br>
                                        Période : Avril - Juillet 2017 [4 mois]
                                        <br>
                                        Centre de formation : INTI Formation – Paris Tour Montparnasse
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Apprentissage du langage informatique Java / Apprentissage de la
                                                conception orientée objet
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Formation Java -->

                        <br>

                        <!-- Diplome ingenieur -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Diplôme d'ingénieur généraliste
                                        <br>
                                        Période : Septembre 2013 - Août 2016 [3 ans]
                                        <br>
                                        Centre de formation : École Supérieure d'Ingénieurs Léonard de Vinci – Paris La
                                        Défense
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Nouvelles énergies / Informatique / Finance / Mécanique numérique
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Diplome ingenieur -->
                    </div>
                </div>
                <!-- Formations -->

                <br>

                <!-- Centre d'intérêts -->
                <div class="card">
                    <div class="card-header">
                        <h4>
                            Centre d'intérêts
                        </h4>
                    </div>
                    <div class="card-body">
                        <ul>
                            <li>
                                Veille technologique / 
                                Brevets d'inventions
                            </li>
                            <li>
                                Droit français / 
                                Droit européen / 
                                Droit international
                            </li>
                            <li>
                                Physique / 
                                Mathématiques / 
                                Finance / 
                                Economie / 
                                Management / 
                                Marketing / 
                                Spiritualité / 
                                Ufologie
                            </li>
                        </ul>
                    </div>
                </div>
                <!-- Centre d'intérêts -->
            </div>
        </div>
    </div>

    <br>

    <!-- Optional JavaScript -->
    <!-- jQuery first, then Popper.js, then Bootstrap JS -->
    <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
        integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
        crossorigin="anonymous"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
        integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
        crossorigin="anonymous"></script>
    <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
        integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
        crossorigin="anonymous"></script>
</body>

</html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de Monsieur ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Société Holomorphe [SIREN : 883632556]'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body, "CV\\CV_De_Mr_.pdf", configuration=config,
                           options=options)

    # ok
    def test_cv_python_outlook_1(self):
        print("test_cv_python_outlook_1")

        # Reporting
        body = """
        <!doctype html>
        <html lang="en">
        
        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">
        
            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">
        
            <title>CV de Monsieur  - Développeur Full-stack / Python / IA</title>
        </head>
        
        <body>
            <br>
            <div class="container">
                <div class="card">
                    <div class="card-header text-center">
                        <h2>
                            Développeur Full-stack / Python / IA
                        </h2>
                    </div>
                    <div class="card-body text-center">
        
                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Identité
                                    </th>
                                    <th scope="col">
                                        Adresse postale
                                    </th>
                                    <th scope="col">
                                        Téléphone
                                    </th>
                                    <th scope="col">
                                        E-mail
                                    </th>
                                    <th scope="col">
                                        Age
                                    </th>
                                    <th scope="col">
                                        Nationalité
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        Monsieur 
                                    </td>
                                    <td>
                                         93800 Epinay-sur-Seine (France)
                                    </td>
                                    <td>
                                        
                                    </td>
                                    <td>
                                        @outlook.fr
                                    </td>
                                    <td>
                                        28 ans
                                    </td>
                                    <td>
                                        Française
                                    </td>
                                </tr>
                            </tbody>
                        </table>
        
                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Comp&eacute;tences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Front-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Front-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Langage de balisage HTML / 
                                                        Langage CSS /
                                                        Framework Django / 
                                                        Framework MJML pour le responsive email
                                                    </li>
                                                    <li>
                                                        Logiciel Visual Studio Code /
                                                        Librairie Bootstrap / 
                                                        Librairie Leaflet.js
                                                    </li>
                                                    <li>
                                                        Librairie Vis.js / 
                                                        Librairie Artyom.js /
                                                        Librairie Flowchart.js /
                                                        Librairie Babylon.js
                                                    </li>
                                                    <li> 
                                                        Serveur Apache /
                                                        Serveur Nginx
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Front-end -->
        
                                <br>
        
                                <!-- Back-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Back-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Langage Python /
                                                        Framework Flask /
                                                        Librairie BeautifulSoup /
                                                        Librairie pdfkit /
                                                        Librairie xlswriter /
                                                        Librairie xlrd /
                                                        Librairie pywinauto /
                                                        Librairie Selenium WebDriver / 
                                                        Package 'unittests' /
                                                        Logiciel SonarQube / 
                                                        Docker
                                                    </li>
                                                    <li>
                                                        Serveur de bases de données relationnelles MySQL / 
                                                        Serveur de bases de données documents MongoDB /
                                                        Serveur de bases de données graphe de Neo4j
                                                    </li>
                                                    <li>
                                                        Exploration de données / 
                                                        Automatisation de processus robotisés /
                                                        Intégration d'API
                                                    </li>
                                                    <li>
                                                        Carte embarquée Raspberry Pi / 
                                                        Système d’exploitation Raspbian
                                                    </li>
                                                    <li>
                                                        Cartes embarquées Arduino Uno et Arduino Yun /
                                                        Système d’exploitation OpenWRT
                                                    </li>
                                                    <li>
                                                        Logiciel de gestion de version Git / 
                                                        Gestion de mon compte GitHub : https://github.com/Jay4C
                                                    </li>       
                                                    <li>
                                                        FreeCAD / 
                                                        Skidl
                                                    </li> 
                                                    <li>
                                                        Réseaux et systèmes / 
                                                        Protocoles informatiques [HTTP, HTTPS, TCP, IPv4, IPv6, SSH, 
                                                        FTP, SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, Modbus RTU]
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Back-end -->
        
                                <br>
        
                                <!-- Management -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Management
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Brainstorming / 
                                                        Logiciel FreeMind / 
                                                        Logiciel Pencil / 
                                                        Diagrammes UML / 
                                                        Logiciel ArgoUML
                                                    </li>
                                                    <li>
                                                        Logiciel GanttProject / 
                                                        Planification des étapes d'un projet
                                                    </li>
                                                    <li>
                                                        Management de projet / 
                                                        Management des opérations / 
                                                        Management Agile / 
                                                        Management de la qualité
                                                    </li>
                                                    <li>
                                                        Analyse des besoins du client / 
                                                        Établissement d’un cahier des charges
                                                    </li>
                                                    <li>
                                                        Rédaction d’un support technique / 
                                                        Rédaction d'une base de données de connaissances
                                                    </li>
                                                    <li>
                                                        Droit français : "www.legifrance.gouv.fr" /
                                                        Droit européen : "https://europa.eu/european-union/law/find-legislation_fr" /
                                                        Droit international : "www.lexadin.nl"
                                                    </li>
                                                    <li>
                                                        Anglais écrit et oral [TOEIC : 840 points ; Date d’expiration 
                                                        : 11 Juillet 2020]
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Management -->
                            </div>
                        </div>
                        <!-- Compétences -->
        
                        <br>
        
                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Holomorphe -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Gestion globale de l'entreprise
                                                <br>
                                                Période : Mai 2020 – En cours
                                                <br>
                                                Statut professionnel : Mandataire social
                                                <br>
                                                Société : Holomorphe - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'applications Python pour créer des pièces 
                                                        mécaniques de précision pour les technologies de l'hydrogène 
                                                        avec le logiciel FreeCAD disponibles sur 
                                                        https://github.com/Jay4C/Python-Macros-For-FreeCAD
                                                    </li>
                                                    <li>
                                                        Développement d'applications Python pour l'automatisation de 
                                                        processus robotisés disponibles sur 
                                                        https://github.com/Jay4C/Web-Automation
                                                    </li>
                                                    <li>
                                                        Développement d'une application Python pour récupérer toutes 
                                                        les parcelles convenables pour injecter du méthane de synthèse 
                                                        dans un réseau de gaz naturel en France
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Holomorphe -->
                                
                                <br>
                                
                                <!-- Société PASS Technologie -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur informatique polyvalent
                                                <br>
                                                Période : Février 2019 - Avril 2019 [3 mois]
                                                <br>
                                                Statut professionnel : Salarié
                                                <br>
                                                Société : PASS Technologie - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application pour le BIM (Building 
                                                        Information Modeling)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Pass Tech -->
        
                                <br>
        
                                <!-- Kriir -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Ingénieur d'études et développement
                                                <br>
                                                Période : Septembre – Novembre 2018 [3 mois]
                                                <br>
                                                Statut professionnel : Salarié
                                                <br>
                                                Société : Kriir - Ermont
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d’une application web en Java pour utiliser la 
                                                        signature électronique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Kriir -->
        
                                <br>
        
                                <!-- Refuel -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Lecture et écriture sur un automate en modbus
                                                <br>
                                                Période : Juin - Juillet 2018 [2 mois]
                                                <br>
                                                Statut professionnel : Indépendant
                                                <br>
                                                Société : Refuel S.A.S. - Trappes
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de plusieurs fonctionnalités permettant de lire et
                                                        d’écrire sur un automate programmable via le
                                                        protocole Modbus TCP avec le langage Java en liaison WiFi
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Refuel  -->
        
                                <br>
        
                                <!-- SETEC SMART EFFICIENCY -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Génération automatique de rapport énergétique
                                                <br>
                                                Période : Novembre - Décembre 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Indépendant
                                                <br>
                                                Société : Setec Smart Efficiency - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développer un programme Java pour générer un rapport énergétique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SETEC SMART EFFICIENCY -->
        
                                <br>
        
                                <!--  SOCIÉTÉ GAMELOFT -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Vérification de la pertinence des données
                                                <br>
                                                Période : Juin - Juillet 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Indépendant
                                                <br>
                                                Société : Gameloft - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d’une macro Visual Basic pour un fichier Excel 
                                                        permettant de contrôler la cohérence des données saisies dans 
                                                        plusieurs fichiers Excel et consolider ces données dans un seul 
                                                        fichier Excel pour les ressources humaines
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SOCIÉTÉ GAMELOFT -->
                            </div>
                        </div>
                        <!-- Expériences -->
        
                        <br>
        
                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : Septembre 2013 - Août 2016 [3 ans]
                                                <br>
                                                Centre de formation : École Supérieure d'Ingénieurs Léonard de Vinci – 
                                                Paris La Défense
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / 
                                                        Informatique / 
                                                        Finance / 
                                                        Mécanique numérique
                                                    </li>
                                                    <li>
                                                        Veille technologique /
                                                        Physique / 
                                                        Mathématiques /
                                                        Economie / 
                                                        Management / 
                                                        Marketing
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->
                    </div>
                </div>
            </div>
        
            <br>
        
            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>
        
        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de Monsieur ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Développeur Full-stack / Python / IA'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body,
                           "CV\\CV_De__[Developpeur_Python].pdf",
                           configuration=config,
                           options=options
                           )

        """

        """

    # ok
    def test_cv_python_outlook_1_1(self):
        print("test_cv_python_outlook_1_1")

        # Reporting
        body = """
        <!doctype html>
        <html lang="en">

        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

            <title>CV de Monsieur  - Développeur Python</title>
        </head>

        <body>
            <br>
            <div class="container">
                <div class="card">
                    <div class="card-header text-center">
                        <h2>
                            Développeur Python
                        </h2>
                    </div>
                    <div class="card-body text-center">

                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Identité
                                    </th>
                                    <th scope="col">
                                        Lieu de travail actuel
                                    </th>
                                    <th scope="col">
                                        Téléphone
                                    </th>
                                    <th scope="col">
                                        Email
                                    </th>
                                    <th scope="col">
                                        Age
                                    </th>
                                    <th scope="col">
                                        Nationalité
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        Monsieur 
                                    </td>
                                    <td>
                                        Holomorphe / 31 Avenue de Ségur 75007 Paris
                                    </td>
                                    <td>
                                        07.45.75.27.00
                                    </td>
                                    <td>
                                        @outlook.fr
                                    </td>
                                    <td>
                                        28 ans
                                    </td>
                                    <td>
                                        Française
                                    </td>
                                </tr>
                            </tbody>
                        </table>

                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Comp&eacute;tences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Front-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Front-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Django web framework / 
                                                        Langage de balisage HTML / 
                                                        Langage CSS /
                                                        Langage de programmation JavaScript
                                                    </li>
                                                    <li>
                                                        Librairie Bootstrap / 
                                                        Librairie Leaflet.js /
                                                        Librairie Vis.js / 
                                                        Librairie Artyom.js /
                                                        Librairie Flowchart.js /
                                                        Librairie Babylon.js / 
                                                        Librairie Chart.js /
                                                        Librairie MathJax
                                                    </li>
                                                    <li>
                                                        Serveur Apache /
                                                        Serveur Nginx
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Front-end -->

                                <br>

                                <!-- Back-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Back-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Langage de programmation Python /
                                                        Framework Flask /
                                                        Librairie SQLAlchemy /
                                                        Librairie unittests /
                                                        Librairie BeautifulSoup /
                                                        Librairie Selenium WebDriver /
                                                        Librairie pymysql.cursors /
                                                        Librairie requests /
                                                        Librairie pdfkit /
                                                        Librairie xlswriter /
                                                        Librairie xlrd /
                                                        Librairie pywinauto /
                                                        Librairie Skidl /
                                                        Librairie PyPDF2 /
                                                        Librairie reportlab /
                                                        Librairie pylatex / 
                                                        Librairie sympy /
                                                        Librairie numpy /
                                                        Librairie pandas /
                                                        Librairie matplotlib /
                                                        Librairie ahkab /
                                                        Librairie email /
                                                        Librairie smtplib /
                                                        Librairie schemdraw /
                                                        Librairie paramiko /
                                                        Librairie imaplib /
                                                        Librairie imap_tools / 
                                                        Librairie minimalmodbus /
                                                        Librairie pyModbusTCP
                                                    </li>
                                                    <li>
                                                        Serveur de bases de données relationnelles MySQL / 
                                                        Serveur de bases de données documents MongoDB /
                                                        Serveur de bases de données graphe de Neo4j
                                                    </li>
                                                    <li>
                                                        Exploration de données / 
                                                        Automatisation des processus robotisés
                                                    </li>
                                                    <li>
                                                        Carte embarquée Raspberry Pi / 
                                                        Système d’exploitation Raspbian
                                                    </li>
                                                    <li>
                                                        API (Programmableweb) /
                                                        Logiciel de gestion de version Git / 
                                                        Gestion de mon compte GitHub : https://github.com/Jay4C / 
                                                        Docker
                                                    </li>
                                                    <li>
                                                        Logiciel FreeCAD / 
                                                        Logiciel KiCAD / 
                                                        Logiciel SonarQube
                                                    </li>
                                                    <li>
                                                        Windows / 
                                                        Linux /
                                                        Ubuntu /
                                                        Bash /
                                                        Serveur privé virtuel
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Back-end -->

                                <br>

                                <!-- Management -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Management
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Brainstorming / 
                                                        Logiciel FreeMind / 
                                                        Logiciel Pencil / 
                                                        Diagrammes UML / 
                                                        Logiciel ArgoUML
                                                    </li>
                                                    <li>
                                                        Diagramme de Gantt /
                                                        Logiciel GanttProject / 
                                                        Planification des étapes d'un projet
                                                    </li>
                                                    <li>
                                                        Management de la connaissance / 
                                                        Management de projet / 
                                                        Management des opérations / 
                                                        Management du changement / 
                                                        Management de la qualité
                                                    </li>
                                                    <li>
                                                        Analyse des besoins du client / 
                                                        Établissement d’un cahier des charges
                                                    </li>
                                                    <li>
                                                        Rédaction d’un support technique / 
                                                        Rédaction d'une base de données de connaissances
                                                    </li>
                                                    <li>
                                                        Droit français : "www.legifrance.gouv.fr" /
                                                        Droit européen : "https://europa.eu/european-union/law/find-legislation_fr" /
                                                        Droit international : "www.lexadin.nl"
                                                    </li>
                                                    <li>
                                                        Anglais écrit et oral [TOEIC : 840 points ; Date d’expiration 
                                                        : 11 Juillet 2020]
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Management -->
                            </div>
                        </div>
                        <!-- Compétences -->

                        <br>

                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Holomorphe -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur Python
                                                <br>
                                                Période : Mai 2019 – En cours
                                                <br>
                                                Statut professionnel : Président
                                                <br>
                                                Société : Holomorphe - SIREN : 883 632 556 - Paris (75007)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'un site internet pour présenter la société 
                                                        disponible sur 
                                                        https://github.com/Jay4C/Holomorphe_Company/tree/main/holomorphefrontend
                                                        (Python / Django web framework)
                                                    </li>
                                                    <li>
                                                        Développement d'applications Python pour créer des pièces 
                                                        mécaniques de précision pour les technologies de l'hydrogène 
                                                        avec le logiciel FreeCAD disponibles sur 
                                                        https://github.com/Jay4C/Python-Macros-For-FreeCAD (Python)
                                                    </li>
                                                    <li>
                                                        Développement d'applications Python pour l'automatisation de 
                                                        processus robotisés disponibles sur 
                                                        https://github.com/Jay4C/Web-Automation (Python / Selenium)
                                                    </li>
                                                    <li>
                                                        Développement d'une application Python pour récupérer toutes 
                                                        les parcelles convenables pour injecter du méthane de synthèse 
                                                        dans un réseau de gaz naturel en France (Python / API Open data)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour programmer des robots 
                                                        d'exploration de données disponibles sur 
                                                        https://github.com/Jay4C/Web-Scraping (Python / BeautifulSoup)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour faire du mailing 
                                                        (Python / smtplib)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour préparer des requêtes 
                                                        vers des API disponibles sur https://github.com/Jay4C/API 
                                                        (Python / API)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour faire des diagrammes de 
                                                        circuits électroniques disponibles sur 
                                                        https://github.com/Jay4C/Electronic_Circuit_Diagram_With_Python_Schemdraw
                                                        (Python / Librairie schemdraw)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour fabriquer des 
                                                        circuits électroniques disponibles sur 
                                                        https://github.com/Jay4C/Electronic_Design_Automation_With_Python_Skidl 
                                                        (Python / Librairie Skidl)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour simuler des 
                                                        circuits électroniques disponibles sur 
                                                        https://github.com/Jay4C/Electronic_Circuit_Simulation_With_Python
                                                        (Python / Librairie ahkab / Librairie pyspice)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Holomorphe -->

                                <br>

                                <!-- Entreprise ALOYAU -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur Python
                                                <br>
                                                Période : Octobre 2016  - Avril 2019 [30 mois]
                                                <br>
                                                Statut professionnel : Auto-entrepreneur
                                                <br>
                                                Société : Entreprise ALOYAU - SIREN : 823 502 222 - Saint-Ouen (93400)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de tests unitaires pour paramétrer des 
                                                        trajectoires de bras-robot avec la carte Raspberry Pi (Python)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour concevoir des circuits 
                                                        électroniques sous le logiciel KiCAD (Python)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour convertir des séries 
                                                        d'images en vidéos (Python)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Entreprise ALOYAU -->
                                
                                <br>

                                <!-- Société Agronergy -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur Python
                                                <br>
                                                Périodes : Avril 2015 - Juillet 2015 [4 mois] 
                                                et Fevrier 2016 - Juillet 2016 [6 mois]
                                                <br>
                                                Statut professionnel : Stagiaire
                                                <br>
                                                Société : Société Agronergy - SIREN : 792 570 137 - Paris (75018)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de programmes informatiques pour collecter 
                                                        des données énergétiques sur des chaudières avec la carte 
                                                        Raspberry Pi grâce au protocole Modbus (Python - Librairie 
                                                        minimalmodbus)
                                                    </li>
                                                    <li>
                                                        Développement de programmes informatiques pour piloter 
                                                        des chaudières à distance avec la carte Raspberry Pi 
                                                        (Python - Librairie pyModbusTCP)
                                                    </li>
                                                    <li>
                                                        Développement de programmes informatiques pour stocker des 
                                                        données énergétiques issues de chaudières dans un serveur de 
                                                        base de données MySQL (Python)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Agronergy -->
                            </div>
                        </div>
                        <!-- Expériences -->

                        <br>

                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : Septembre 2013 - Août 2016 [3 ans]
                                                <br>
                                                Centre de formation : École Supérieure d'Ingénieurs Léonard de Vinci – 
                                                Paris La Défense
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / 
                                                        Informatique / 
                                                        Finance / 
                                                        Mécanique numérique
                                                    </li>
                                                    <li>
                                                        Veille technologique /
                                                        Physique / 
                                                        Mathématiques /
                                                        Economie / 
                                                        Management / 
                                                        Marketing /
                                                        Comptabilité
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->
                    </div>
                </div>
            </div>

            <br>

            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>

        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de Monsieur ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Développeur Python'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body,
                           "CV\\CV_De__[Developpeur_Python].pdf",
                           configuration=config,
                           options=options
                           )

    # ok
    def test_cv_python_outlook_1_1_1(self):
        print("test_cv_python_outlook_1_1_1")

        # Reporting
        body = """
        <!doctype html>
        <html lang="en">

        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

            <title>CV de Monsieur  - Développeur Python</title>
        </head>

        <body>
            <br>
            <div class="container">
                <div class="card">
                    <div class="card-header text-center">
                        <h2>
                            Développeur Python
                        </h2>
                    </div>
                    <div class="card-body text-center">

                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Identité
                                    </th>
                                    <th scope="col">
                                        Adresse postale
                                    </th>
                                    <th scope="col">
                                        Téléphone
                                    </th>
                                    <th scope="col">
                                        Email
                                    </th>
                                    <th scope="col">
                                        Age
                                    </th>
                                    <th scope="col">
                                        Nationalité
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        Monsieur 
                                    </td>
                                    <td>
                                         93800 Epinay-sur-Seine
                                    </td>
                                    <td>
                                        07.45.75.27.00
                                    </td>
                                    <td>
                                        @outlook.fr
                                    </td>
                                    <td>
                                        28 ans
                                    </td>
                                    <td>
                                        Française
                                    </td>
                                </tr>
                            </tbody>
                        </table>

                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Comp&eacute;tences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Front-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Front-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Django web framework / 
                                                        Langage de balisage HTML / 
                                                        Langage CSS /
                                                        Langage de programmation JavaScript
                                                    </li>
                                                    <li>
                                                        Librairie Bootstrap / 
                                                        Librairie Leaflet.js /
                                                        Librairie Vis.js / 
                                                        Librairie Artyom.js /
                                                        Librairie Flowchart.js /
                                                        Librairie Babylon.js / 
                                                        Librairie Chart.js /
                                                        Librairie MathJax
                                                    </li>
                                                    <li>
                                                        Serveur Apache /
                                                        Serveur Nginx
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Front-end -->

                                <br>

                                <!-- Back-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Back-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Langage de programmation Python /
                                                        Framework Flask /
                                                        Librairie SQLAlchemy /
                                                        Librairie unittests /
                                                        Librairie BeautifulSoup /
                                                        Librairie Selenium WebDriver /
                                                        Librairie pymysql.cursors /
                                                        Librairie requests /
                                                        Librairie pdfkit /
                                                        Librairie xlswriter /
                                                        Librairie xlrd /
                                                        Librairie pywinauto /
                                                        Librairie Skidl /
                                                        Librairie PyPDF2 /
                                                        Librairie reportlab /
                                                        Librairie pylatex / 
                                                        Librairie sympy /
                                                        Librairie numpy /
                                                        Librairie pandas /
                                                        Librairie matplotlib /
                                                        Librairie ahkab /
                                                        Librairie email /
                                                        Librairie smtplib /
                                                        Librairie schemdraw /
                                                        Librairie paramiko /
                                                        Librairie imaplib /
                                                        Librairie imap_tools / 
                                                        Librairie minimalmodbus /
                                                        Librairie pyModbusTCP
                                                    </li>
                                                    <li>
                                                        Serveur de bases de données relationnelles MySQL / 
                                                        Serveur de bases de données documents MongoDB /
                                                        Serveur de bases de données graphe de Neo4j
                                                    </li>
                                                    <li>
                                                        Exploration de données / 
                                                        Automatisation des processus robotisés
                                                    </li>
                                                    <li>
                                                        Carte embarquée Raspberry Pi / 
                                                        Système d’exploitation Raspbian
                                                    </li>
                                                    <li>
                                                        Logiciel de gestion de version Git / 
                                                        Gestion de mon compte GitHub : https://github.com/Jay4C / 
                                                        Docker
                                                    </li>
                                                    <li>
                                                        Logiciel FreeCAD / 
                                                        Logiciel KiCAD / 
                                                        Logiciel SonarQube
                                                    </li>
                                                    <li>
                                                        Windows / 
                                                        Linux /
                                                        Ubuntu /
                                                        Bash /
                                                        Serveur privé virtuel
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Back-end -->

                                <br>

                                <!-- Management -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Management
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Brainstorming / 
                                                        Logiciel FreeMind / 
                                                        Logiciel Pencil / 
                                                        Diagrammes UML / 
                                                        Logiciel Gaphor
                                                    </li>
                                                    <li>
                                                        Diagramme de Gantt /
                                                        Logiciel GanttProject / 
                                                        Planification des étapes d'un projet
                                                    </li>
                                                    <li>
                                                        Management de la connaissance / 
                                                        Management de projet / 
                                                        Management des opérations / 
                                                        Management du changement / 
                                                        Management de la qualité
                                                    </li>
                                                    <li>
                                                        Analyse des besoins du client / 
                                                        Établissement d’un cahier des charges
                                                    </li>
                                                    <li>
                                                        Rédaction d’un support technique / 
                                                        Rédaction d'une base de données de connaissances
                                                    </li>
                                                    <li>
                                                        Droit français : "www.legifrance.gouv.fr" /
                                                        Droit européen : "https://europa.eu/european-union/law/find-legislation_fr" /
                                                        Droit international : "www.lexadin.nl"
                                                    </li>
                                                    <li>
                                                        Anglais écrit et oral [TOEIC : 840 points ; Date d’expiration 
                                                        : 11 Juillet 2020]
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Management -->
                            </div>
                        </div>
                        <!-- Compétences -->

                        <br>

                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Holomorphe -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur Python
                                                <br>
                                                Période : Mai 2019 – En cours
                                                <br>
                                                Statut professionnel : Président
                                                <br>
                                                Société : Holomorphe - SIREN : 883 632 556 - Paris (75007)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'un site internet pour présenter la société 
                                                        disponible sur 
                                                        https://github.com/Jay4C/Holomorphe_Company/tree/main/holomorphefrontend
                                                        (Python / Django web framework)
                                                    </li>
                                                    <li>
                                                        Développement d'un programme d'API disponible sur 
                                                        https://github.com/Jay4C/Holomorphe_Company/tree/main/holomorphebackend
                                                        (Python / Flask API)
                                                    </li>
                                                    <li>
                                                        Développement d'applications Python pour créer des pièces 
                                                        mécaniques de précision pour les technologies de l'hydrogène 
                                                        avec le logiciel FreeCAD disponibles sur 
                                                        https://github.com/Jay4C/Python-Macros-For-FreeCAD (Python)
                                                    </li>
                                                    <li>
                                                        Développement d'applications Python pour l'automatisation de 
                                                        processus robotisés disponibles sur 
                                                        https://github.com/Jay4C/Web-Automation (Python / Selenium)
                                                    </li>
                                                    <li>
                                                        Développement d'une application Python pour récupérer toutes 
                                                        les parcelles convenables pour injecter du méthane de synthèse 
                                                        dans un réseau de gaz naturel en France (Python / API Open data)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour programmer des robots 
                                                        d'exploration de données disponibles sur 
                                                        https://github.com/Jay4C/Web-Scraping (Python / BeautifulSoup)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour faire du mailing 
                                                        (Python / smtplib)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour préparer des requêtes 
                                                        vers des API disponibles sur https://github.com/Jay4C/API 
                                                        (Python / API)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour faire des diagrammes de 
                                                        circuits électroniques disponibles sur 
                                                        https://github.com/Jay4C/Electronic_Circuit_Diagram_With_Python_Schemdraw
                                                        (Python / Librairie schemdraw)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour fabriquer des 
                                                        circuits électroniques disponibles sur 
                                                        https://github.com/Jay4C/Electronic_Design_Automation_With_Python_Skidl 
                                                        (Python / Librairie Skidl)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour simuler des 
                                                        circuits électroniques disponibles sur 
                                                        https://github.com/Jay4C/Electronic_Circuit_Simulation_With_Python
                                                        (Python / Librairie ahkab / Librairie pyspice)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Holomorphe -->

                                <br>

                                <!-- Entreprise ALOYAU -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur Python
                                                <br>
                                                Période : Octobre 2016  - Avril 2019 [30 mois]
                                                <br>
                                                Statut professionnel : Auto-entrepreneur
                                                <br>
                                                Société : Entreprise ALOYAU - SIREN : 823 502 222 - Saint-ouen (93400)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de tests unitaires pour paramétrer des 
                                                        trajectoires de bras-robot avec la carte Raspberry Pi (Python)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour concevoir des circuits 
                                                        électroniques sous le logiciel KiCAD orientés pour l'agriculture 
                                                        (Python)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour convertir des séries 
                                                        d'images en vidéos (Python)
                                                    </li>
                                                    <li>
                                                        Gestion de visite guidée de biens immobiliers pour le 
                                                        compte de la société Flatsy en utilisant une application mobile
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Entreprise ALOYAU -->

                                <br>

                                <!-- Société Agronergy -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur Python
                                                <br>
                                                Périodes : Avril 2015 - Juillet 2015 [4 mois] 
                                                et Fevrier 2016 - Juillet 2016 [6 mois]
                                                <br>
                                                Statut professionnel : Stagiaire
                                                <br>
                                                Société : Société Agronergy - SIREN : 792 570 137 - Paris (75018)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de programmes informatiques pour collecter 
                                                        des données énergétiques sur des chaudières avec la carte 
                                                        Raspberry Pi grâce au protocole Modbus (Python - Librairie 
                                                        minimalmodbus)
                                                    </li>
                                                    <li>
                                                        Développement de programmes informatiques pour piloter 
                                                        des chaudières à distance avec la carte Raspberry Pi 
                                                        (Python - Librairie pyModbusTCP)
                                                    </li>
                                                    <li>
                                                        Développement de programmes informatiques pour stocker des 
                                                        données énergétiques issues de chaudières dans un serveur de 
                                                        base de données MySQL (Python)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Agronergy -->
                            </div>
                        </div>
                        <!-- Expériences -->

                        <br>

                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : Septembre 2013 - Août 2016 [3 ans]
                                                <br>
                                                Centre de formation : École Supérieure d'Ingénieurs Léonard de Vinci – 
                                                Paris La Défense
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / 
                                                        Informatique / 
                                                        Finance / 
                                                        Mécanique numérique
                                                    </li>
                                                    <li>
                                                        Veille technologique /
                                                        Physique / 
                                                        Mathématiques /
                                                        Economie / 
                                                        Management / 
                                                        Marketing /
                                                        Comptabilité
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                                
                                <br>
                                
                                <!-- CPGE PTSI-PT -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Classes préparatoires aux grandes écoles PTSI-PT
                                                <br>
                                                Période : Août 2011 - Août 2013 [2 ans]
                                                <br>
                                                Centre de formation : Lycée Lislet Geofrroy à l'Ile de La Réunion
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Mathématiques / 
                                                        Chimie / 
                                                        Physique / 
                                                        Technologie /
                                                        Sciences de l'ingénieur
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- CPGE PTSI-PT -->
                                
                                <br>
                                
                                <!-- Baccalauréat scientifique -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Baccalauréat scientifique
                                                <br>
                                                Période : Août 2010 - Août 2011 [1 an]
                                                <br>
                                                Centre de formation : Lycée Roland Garros à l'Ile de La Réunion
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Sciences de l'ingénieur / 
                                                        Mathématiques / 
                                                        Physique / 
                                                        Chimie / 
                                                        Anglais / 
                                                        Français / 
                                                        Histoire
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Baccalauréat scientifique -->
                            </div>
                        </div>
                        <!-- Formations -->
                    </div>
                </div>
            </div>

            <br>

            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>

        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de Monsieur ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Développeur Python'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body,
                           "CV\\CV__[Developpeur_Python].pdf",
                           configuration=config,
                           options=options
                           )

    # ok
    def test_cv_python_outlook_1_1_anonymise(self):
        print("test_cv_python_outlook_1_1_anonymise")

        # Reporting
        body = """
        <!doctype html>
        <html lang="en">

        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

            <title>CV de Monsieur  - Développeur Python</title>
        </head>

        <body>
            <br>
            <div class="container">
                <div class="card">
                    <div class="card-header text-center">
                        <h2>
                            Développeur Python
                        </h2>
                    </div>
                    <div class="card-body text-center">
                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Comp&eacute;tences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Front-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Front-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Django web framework / 
                                                        Langage de balisage HTML / 
                                                        Langage CSS
                                                    </li>
                                                    <li>
                                                        Librairie Bootstrap / 
                                                        Librairie Leaflet.js /
                                                        Librairie Vis.js / 
                                                        Librairie Artyom.js /
                                                        Librairie Flowchart.js /
                                                        Librairie Babylon.js / 
                                                        Librairie Chart.js /
                                                        Librairie MathJax
                                                    </li>
                                                    <li>
                                                        Serveur Apache /
                                                        Serveur Nginx
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Front-end -->

                                <br>

                                <!-- Back-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Back-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Langage de programmation Python /
                                                        Framework Flask /
                                                        Librairie SQLAlchemy /
                                                        Librairie unittests /
                                                        Librairie BeautifulSoup /
                                                        Librairie Selenium WebDriver /
                                                        Librairie pymysql.cursors /
                                                        Librairie requests /
                                                        Librairie pdfkit /
                                                        Librairie xlswriter /
                                                        Librairie xlrd /
                                                        Librairie pywinauto /
                                                        Librairie Skidl /
                                                        Librairie PyPDF2 /
                                                        Librairie reportlab /
                                                        Librairie pylatex / 
                                                        Librairie sympy /
                                                        Librairie numpy /
                                                        Librairie pandas /
                                                        Librairie matplotlib /
                                                        Librairie ahkab /
                                                        Librairie email /
                                                        Librairie smtplib /
                                                        Librairie schemdraw /
                                                        Librairie paramiko /
                                                        Librairie imaplib /
                                                        Librairie imap_tools / 
                                                        Librairie minimalmodbus /
                                                        Librairie pyModbusTCP
                                                    </li>
                                                    <li>
                                                        Serveur de bases de données relationnelles MySQL / 
                                                        Serveur de bases de données documents MongoDB /
                                                        Serveur de bases de données graphe de Neo4j
                                                    </li>
                                                    <li>
                                                        Exploration de données / 
                                                        Automatisation des processus robotisés
                                                    </li>
                                                    <li>
                                                        Carte embarquée Raspberry Pi / 
                                                        Système d’exploitation Raspbian
                                                    </li>
                                                    <li>
                                                        API (Programmableweb) /
                                                        Logiciel de gestion de version Git / 
                                                        GitHub
                                                        Docker
                                                    </li>
                                                    <li>
                                                        Logiciel FreeCAD / 
                                                        Logiciel KiCAD / 
                                                        Logiciel SonarQube
                                                    </li>
                                                    <li>
                                                        Windows / 
                                                        Linux /
                                                        Ubuntu /
                                                        Bash /
                                                        Serveur privé virtuel
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Back-end -->

                                <br>

                                <!-- Management -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Management
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Brainstorming / 
                                                        Logiciel FreeMind / 
                                                        Logiciel Pencil / 
                                                        Diagrammes UML / 
                                                        Logiciel ArgoUML
                                                    </li>
                                                    <li>
                                                        Diagramme de Gantt /
                                                        Logiciel GanttProject / 
                                                        Planification des étapes d'un projet
                                                    </li>
                                                    <li>
                                                        Management de la connaissance / 
                                                        Management de projet / 
                                                        Management des opérations / 
                                                        Management du changement / 
                                                        Management de la qualité
                                                    </li>
                                                    <li>
                                                        Analyse des besoins du client / 
                                                        Établissement d’un cahier des charges
                                                    </li>
                                                    <li>
                                                        Rédaction d’un support technique / 
                                                        Rédaction d'une base de données de connaissances
                                                    </li>
                                                    <li>
                                                        Droit français : "www.legifrance.gouv.fr" /
                                                        Droit européen : "https://europa.eu/european-union/law/find-legislation_fr" /
                                                        Droit international : "www.lexadin.nl"
                                                    </li>
                                                    <li>
                                                        Anglais écrit et oral [TOEIC : 840 points ; Date d’expiration 
                                                        : 11 Juillet 2020]
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Management -->
                            </div>
                        </div>
                        <!-- Compétences -->

                        <br>

                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Holomorphe -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur Python
                                                <br>
                                                Période : Mai 2019 – En cours
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'un site internet (Python / Django web framework)
                                                    </li>
                                                    <li>
                                                        Développement d'applications Python pour créer des pièces 
                                                        mécaniques de précision pour les technologies de l'hydrogène 
                                                        avec le logiciel FreeCAD (Python)
                                                    </li>
                                                    <li>
                                                        Développement d'applications Python pour l'automatisation de 
                                                        processus robotisés (Python / Selenium)
                                                    </li>
                                                    <li>
                                                        Développement d'une application Python pour récupérer toutes 
                                                        les parcelles convenables pour injecter du méthane de synthèse 
                                                        dans un réseau de gaz naturel en France (Python / 
                                                        API Open data)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour programmer des robots 
                                                        d'exploration de données (Python / BeautifulSoup)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour faire du mailing 
                                                        (Python / smtplib)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour préparer des requêtes 
                                                        vers des API (Python / API)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour faire des diagrammes de 
                                                        circuits électroniques (Python / Librairie schemdraw)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour fabriquer des 
                                                        circuits électroniques (Python / Librairie Skidl)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour simuler des 
                                                        circuits électroniques (Python / Librairie ahkab 
                                                        / Librairie pyspice)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Holomorphe -->

                                <br>

                                <!-- Entreprise ALOYAU -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur Python
                                                <br>
                                                Période : Octobre 2016  - Avril 2019 [30 mois]
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de tests unitaires pour paramétrer des 
                                                        trajectoires de bras-robot avec la carte Raspberry Pi (Python)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour concevoir des circuits 
                                                        électroniques sous le logiciel KiCAD (Python)
                                                    </li>
                                                    <li>
                                                        Développement de tests unitaires pour converir des séries 
                                                        d'images en vidéos (Python)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Entreprise ALOYAU -->

                                <br>

                                <!-- Société Agronergy -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur Python
                                                <br>
                                                Périodes : Avril 2015 - Juillet 2015 [4 mois] 
                                                et Fevrier 2016 - Juillet 2016 [6 mois]
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de programmes informatiques pour collecter 
                                                        des données énergétiques sur des chaudières avec la carte 
                                                        Raspberry Pi grâce au protocole Modbus (Python - Librairie 
                                                        minimalmodbus)
                                                    </li>
                                                    <li>
                                                        Développement de programmes informatiques pour piloter 
                                                        des chaudières à distance avec la carte Raspberry Pi 
                                                        (Python - Librairie pyModbusTCP)
                                                    </li>
                                                    <li>
                                                        Développement de programmes informatiques pour stocker des 
                                                        données énergétiques issues de chaudières dans un serveur de 
                                                        base de données MySQL (Python)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Agronergy -->
                            </div>
                        </div>
                        <!-- Expériences -->

                        <br>

                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : Septembre 2013 - Août 2016 [3 ans]
                                                <br>
                                                Centre de formation : École Supérieure d'Ingénieurs Léonard de Vinci – 
                                                Paris La Défense
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / 
                                                        Informatique / 
                                                        Finance / 
                                                        Mécanique numérique
                                                    </li>
                                                    <li>
                                                        Veille technologique /
                                                        Physique / 
                                                        Mathématiques /
                                                        Economie / 
                                                        Management / 
                                                        Marketing /
                                                        Comptabilité
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->
                    </div>
                </div>
            </div>

            <br>

            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>

        </html>
        """

        options = {
            'page-size': 'A4',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Développeur Python'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body,
                           "CV\\CV_J_A_[Developpeur_Python].pdf",
                           configuration=config,
                           options=options
                           )

    # ok
    def test_cv_python_outlook_2(self):
        print("test_cv_python_outlook_2")

        # Reporting
        body = """
        <!doctype html>
        <html lang="en">
        
        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">
        
            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">
        
            <title>CV de Monsieur  - Développeur Python / RPA / Data miner / Full-stack</title>
        </head>
        
        <body>
            <br>
            <div class="container">
                <div class="card">
                    <div class="card-header text-center">
                        <h2>
                            Développeur Python / RPA / Data miner / Full-stack
                        </h2>
                    </div>
                    <div class="card-body text-center">
        
                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Identité
                                    </th>
                                    <th scope="col">
                                        Adresse postale
                                    </th>
                                    <th scope="col">
                                        Téléphone
                                    </th>
                                    <th scope="col">
                                        E-mail
                                    </th>
                                    <th scope="col">
                                        Age
                                    </th>
                                    <th scope="col">
                                        Nationalité
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        Monsieur 
                                    </td>
                                    <td>
                                         93800 Epinay-sur-Seine (France)
                                    </td>
                                    <td>
                                        
                                    </td>
                                    <td>
                                        aymerick.aloyau@outlook.fr
                                    </td>
                                    <td>
                                        28 ans
                                    </td>
                                    <td>
                                        Française
                                    </td>
                                </tr>
                            </tbody>
                        </table>
        
                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Comp&eacute;tences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Anglais -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Anglais
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Anglais écrit et oral [TOEIC : 840 points ; Date d’expiration : 11 Juillet 2020]
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Anglais -->
        
                                <br>
        
                                <!-- Langages de programmation de prédilection -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Langages de programmation de prédilection
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>Langage Python / Langage Java / Langage JavaScript</li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Langages de programmation de prédilection -->
        
                                <br>
        
                                <!-- Bureautique -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Bureautique
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>Logiciels Microsoft Excel, Word, PowerPoint, Access, Outlook</li>
                                                    <li>Logiciels Open office Calc, Writer</li>
                                                    <li>
                                                        Logiciel IntelliJ pour Java /
                                                        Logiciel PyCharm pour Python
                                                    </li>
                                                    <li>
                                                        Logiciel Pencil pour faire des schémas simplifiés / 
                                                        Logiciel GanttProject pour le management de projet
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Bureautique -->
        
                                <br>
        
                                <!-- Front-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Front-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Langage de balisage HTML / 
                                                        Langage CSS /
                                                        Framework JQuery / 
                                                        Framework Django / 
                                                        Framework MJML pour le responsive email
                                                    </li>
                                                    <li>
                                                        Logiciel Visual Studio Code /
                                                        Framework Angular 7 /
                                                        Librairie Bootstrap / 
                                                        Librairie Leaflet.js
                                                    </li>
                                                    <li>
                                                        Librairie Vis.js / 
                                                        Librairie Artyom.js /
                                                        Librairie Flowchart.js
                                                    </li>
                                                    <li>
                                                        Serveur Tomcat / 
                                                        Serveur Apache /
                                                        Serveur Nginx
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Front-end -->
        
                                <br>
        
                                <!-- Back-end -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Back-end
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Langage Python /
                                                        Framework Flask /
                                                        Librairie BeautifulSoup /
                                                        Librairie pdfkit /
                                                        Librairie xlswriter /
                                                        Librairie xlrd /
                                                        Librairie pywinauto /
                                                        Librairie Hyperledger Sawtooth
                                                    </li>
                                                    <li>
                                                        Langage Java /
                                                        Gestionnaire de dépendances Maven / 
                                                        Gestionnaire de dépendances Gradle /
                                                        Librairie Spring /
                                                        Librairie Jsoup /
                                                        Librairie Selenium WebDriver
                                                    </li>
                                                    <li>
                                                        Langage JavaScript /
                                                        Langage Visual Basic /
                                                        Langage C# /
                                                        Langage TypeScript
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Back-end -->
        
                                <br>
        
                                <!-- Tests unitaires et qualité -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Tests unitaires et qualité
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Package Unittests avec Python / 
                                                        Package JUnit avec Java / 
                                                        Mockito avec Java / 
                                                        WireMock avec Java
                                                    </li>
                                                    <li>
                                                        Logiciel SonarQube
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Tests unitaires et qualité -->
        
                                <br>
        
                                <!-- Systèmes de bases de données -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Systèmes de bases de données
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Serveur de bases de données MySQL / 
                                                        Logiciel MySQL Workbench
                                                    </li>
                                                    <li>
                                                        Serveur de bases de données graphe de Neo4j / 
                                                        Serveur de bases de données documents MongoDB
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Systèmes de bases de données -->
        
                                <br>
        
                                <!-- Outils de déploiement -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Outils de déploiement
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Docker / Git
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Outils de déploiement -->
        
                                <br>
        
                                <!-- Intelligence artificielle -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Intelligence artificielle
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Exploration de données / 
                                                        Gestion de la connaissance / 
                                                        Automatisation de processus robotisés /
                                                        RPA
                                                    </li>
                                                    <li>
                                                        Cloud computing [Open data; Autodesk Forge; Netheos] / 
                                                        Reconnaissance vocale / 
                                                        Vision par ordinateur
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Intelligence artificielle -->
        
                                <br>
        
                                <!-- Mobile -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Mobile
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Système d'exploitation Android / 
                                                        Logiciel Android Studio
                                                    </li>
                                                    <li>
                                                        Framework Kivy pour le développement mobile multiplateforme développé 
                                                        en Python
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Mobile -->
        
                                <br>
        
                                <!-- Systèmes embarqués -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Systèmes embarqués
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Carte embarquée Raspberry Pi / 
                                                        Système d’exploitation Raspbian
                                                    </li>
                                                    <li>
                                                        Cartes embarquées Arduino Uno et Arduino Yun /
                                                        Système d’exploitation OpenWRT
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Systèmes embarqués -->
        
                                <br>
        
                                <!-- Conception assistée par ordinateur -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Conception assistée par ordinateur
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Logiciel FreeCAD développé en Python pour la conception de pièces 
                                                        mécaniques de précision (https://www.freecadweb.org/?lang=fr)
                                                    </li>
                                                    <li>
                                                        Logiciel Skidl développé en Python pour la génération de fichiers 
                                                        netlist pour la fabrication de circuits imprimés
                                                        (https://devbisme.github.io/skidl/)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Conception assistée par ordinateur -->
        
                                <br>
        
                                <!-- Gestion de version -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Gestion de version
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Logiciel de gestion de version Git / 
                                                        Gestion de mon compte GitHub : https://github.com/Jay4C
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Gestion de version -->
        
                                <br>
        
                                <!-- Réseaux et Systèmes -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Réseaux et systèmes
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Système d’exploitation Windows / 
                                                        Système d’exploitation Linux
                                                    </li>
                                                    <li>
                                                        Protocoles informatiques [HTTP, HTTPS, TCP, IPv4, IPv6, SSH, FTP, 
                                                        SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, Modbus RTU]
                                                    </li>
                                                    <li>
                                                        Librairie MinimalModbus avec Python / 
                                                        Librairie Jlibmodbus avec Java
                                                    </li>
                                                    <li>
                                                        Routeur 4G / 
                                                        Port Forwarding / 
                                                        DynDNS / 
                                                        Dongle 3G
                                                    </li>
                                                    <li>
                                                        Logiciel Wireshark / 
                                                        Logiciel PuTTY / 
                                                        Logiciel WinSCP /
                                                        Logiciel Postman /
                                                        Logiciel TeamViewer / 
                                                        Logiciel VirtualBox /
                                                        Logiciel Tor /
                                                        Logiciel Angry IP Scanner
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Réseaux et Systèmes -->
        
                                <br>
        
                                <!-- Gestion de projet et communications -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Gestion de projet et communications
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Brainstorming / 
                                                        Logiciel FreeMind / 
                                                        Logiciel Pencil / 
                                                        Diagrammes UML / 
                                                        Logiciel ArgoUML
                                                    </li>
                                                    <li>
                                                        Logiciel GanttProject / 
                                                        Planification des étapes d'un projet
                                                    </li>
                                                    <li>
                                                        Management de projet / 
                                                        Management des opérations / 
                                                        Management Agile / 
                                                        Management de la qualité
                                                    </li>
                                                    <li>
                                                        Analyse des besoins du client / 
                                                        Établissement d’un cahier des charges
                                                    </li>
                                                    <li>
                                                        Rédaction d’un support technique / 
                                                        Rédaction d'une base de données de connaissances
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Gestion de projet et communications -->
        
                                <br>
        
                                <!-- Droit -->
                                <table class="table table-bordered">
                                    <thead>
                                        <tr>
                                            <th scope="col">
                                                Droit
                                            </th>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Code général des impôts / 
                                                        Code monétaire et financier / 
                                                        Code de l'environnement / 
                                                        Code de l'énergie
                                                    </li>
                                                    <li>
                                                        Code des relations entre le public et l’administration / 
                                                        Code de la construction et de l'habitation
                                                    </li>
                                                    <li>
                                                        Code de l'urbanisme /
                                                        Code des postes et des communications électroniques /                                                 
                                                    </li>
                                                    <li>
                                                        Site internet "www.legifrance.gouv.fr" /
                                                        Site internet "www.lexadin.nl" /
                                                        Site internet "https://europa.eu/european-union/law/find-legislation_fr"
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Droit -->
                            </div>
                        </div>
                        <!-- Compétences -->
        
                        <br>
        
                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Société PASS Technologie -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Développeur informatique polyvalent
                                                <br>
                                                Période : Février 2019 - Avril 2019 [3 mois]
                                                <br>
                                                Statut professionnel : Salarié
                                                <br>
                                                Société : PASS Technologie - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application pour le BIM (Building Information
                                                        Modeling)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Pass Tech -->
        
                                <br>
        
                                <!-- Kriir -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Ingénieur d'études et développement
                                                <br>
                                                Période : Septembre – Novembre 2018 [3 mois]
                                                <br>
                                                Statut professionnel : Salarié
                                                <br>
                                                Société : Kriir - Ermont
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d’une application web en Java pour utiliser la signature
                                                        électronique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Kriir -->
        
                                <br>
        
                                <!-- Refuel -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Lecture et écriture sur un automate en modbus
                                                <br>
                                                Période : Juin - Juillet 2018 [2 mois]
                                                <br>
                                                Statut professionnel : Indépendant
                                                <br>
                                                Société : Refuel S.A.S. - Trappes
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de plusieurs fonctionnalités permettant de lire et
                                                        d’écrire sur un automate programmable via le
                                                        protocole Modbus TCP avec le langage Java en liaison WiFi
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Refuel  -->
        
                                <br>
        
                                <!-- SETEC SMART EFFICIENCY -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Génération automatique de rapport énergétique
                                                <br>
                                                Période : Novembre - Décembre 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Indépendant
                                                <br>
                                                Société : Setec Smart Efficiency - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développer un programme Java pour générer un rapport énergétique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SETEC SMART EFFICIENCY -->
        
                                <br>
        
                                <!--  SOCIÉTÉ GAMELOFT -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Vérification de la pertinence des données
                                                <br>
                                                Période : Juin - Juillet 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Indépendant
                                                <br>
                                                Société : Gameloft - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d’une macro Visual Basic pour un fichier Excel 
                                                        permettant de contrôler la cohérence des données saisies dans 
                                                        plusieurs fichiers Excel et consolider ces données dans un seul 
                                                        fichier Excel pour les ressources humaines
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SOCIÉTÉ GAMELOFT -->
        
                                <br>
                            </div>
                        </div>
                        <!-- Expériences -->
        
                        <br>
        
                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Formation Java -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Formation Java / JEE
                                                <br>
                                                Période : Avril - Juillet 2017 [4 mois]
                                                <br>
                                                Centre de formation : INTI Formation – Paris Tour Montparnasse
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Apprentissage du langage informatique Java / Apprentissage de la
                                                        conception orientée objet
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Formation Java -->
        
                                <br>
        
                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Intitulé : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : Septembre 2013 - Août 2016 [3 ans]
                                                <br>
                                                Centre de formation : École Supérieure d'Ingénieurs Léonard de Vinci – Paris La
                                                Défense
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / Informatique / Finance / Mécanique numérique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->
        
                        <br>
        
                        <!-- Centre d'intérêts -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Centre d'intérêts
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Veille technologique / 
                                        Brevets d'inventions
                                    </li>
                                    <li>
                                        Droit français / 
                                        Droit européen / 
                                        Droit international
                                    </li>
                                    <li>
                                        Physique / 
                                        Mathématiques / 
                                        Finance / 
                                        Economie / 
                                        Management / 
                                        Marketing / 
                                        Spiritualité / 
                                        Ufologie
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Centre d'intérêts -->
                    </div>
                </div>
            </div>
        
            <br>
        
            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>
        
        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de Monsieur ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Développeur Python / RPA / Data miner / Full-stack'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body,
                           "CV\\CV_De__[Developpeur_Python]_v1.pdf",
                           configuration=config,
                           options=options
                           )

    # ok
    def test_cv_employe_polyvalent(self):
        print("test_cv_employe_polyvalent")

        # Reporting
        body = """
<!doctype html>
<html lang="en">

<head>
    <!-- Required meta tags -->
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

    <!-- Bootstrap CSS -->
    <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
        integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

    <title>CV de Monsieur  - Employé polyvalent / Développeur informatique</title>
</head>

<body>
    <br>
    <div class="container">
        <div class="card">
            <div class="card-header text-center">
                <h2>
                    Employé polyvalent / Développeur informatique
                </h2>
            </div>
            <div class="card-body text-center">

                <!-- Informations personnelles -->
                <table class="table table-striped table-bordered">
                    <thead>
                        <tr>
                            <th scope="col">Identité</th>
                            <th scope="col">Adresse postale</th>
                            <th scope="col">Téléphone</th>
                            <th scope="col">E-mail</th>
                            <th scope="col">&Acirc;ge</th>
                            <th scope="col">Nationalit&eacute;</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr>
                            <td>
                                Monsieur Jason
                                <br>
                                ALOYAU
                            </td>
                            <td>
                                
                                <br>
                                93800 Epinay-sur-Seine
                            </td>
                            <td>
                                
                            </td>
                            <td>
                                
                            </td>
                            <td>
                                28 ans
                            </td>
                            <td>
                                Française
                            </td>
                        </tr>
                    </tbody>
                </table>

                <!--  Compétences -->
                <div class="card">
                    <div class="card-header">
                        <h4>
                            Comp&eacute;tences
                        </h4>
                    </div>
                    <div class="card-body">
                        <!-- Anglais -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Anglais
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Anglais écrit et oral [TOEIC : 840 points ; Date d’expiration : 11 Juillet 2020]
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Anglais -->

                        <br>

                        <!-- Langages de programmation de prédilection -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Langages de programmation de prédilection
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>Langage Python / Langage Java / Langage JavaScript</li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Langages de programmation de prédilection -->

                        <br>

                        <!-- Bureautique -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Bureautique
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>Logiciels Microsoft Excel, Word, PowerPoint, Access, Outlook</li>
                                            <li>Logiciels Open office Calc, Writer</li>
                                            <li>
                                                Logiciel IntelliJ pour Java /
                                                Logiciel PyCharm pour Python
                                            </li>
                                            <li>
                                                Logiciel Pencil pour faire des schémas simplifiés / 
                                                Logiciel GanttProject pour le management de projet
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Bureautique -->

                        <br>

                        <!-- Front-end -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Front-end
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Langage de balisage HTML / 
                                                Langage CSS /
                                                Framework JQuery / 
                                                Framework Django / 
                                                Framework MJML pour le responsive email
                                            </li>
                                            <li>
                                                Logiciel Visual Studio Code /
                                                Framework Angular 7 /
                                                Librairie Bootstrap / 
                                                Librairie Leaflet.js
                                            </li>
                                            <li>
                                                Librairie Vis.js / 
                                                Librairie Artyom.js /
                                                Librairie Flowchart.js
                                            </li>
                                            <li>
                                                Serveur Tomcat / 
                                                Serveur Apache /
                                                Serveur Nginx
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Front-end -->

                        <br>

                        <!-- Back-end -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Back-end
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Langage Python /
                                                Framework Flask /
                                                Librairie BeautifulSoup /
                                                Librairie pdfkit /
                                                Librairie xlswriter /
                                                Librairie xlrd /
                                                Librairie pywinauto /
                                                Librairie Hyperledger Sawtooth
                                            </li>
                                            <li>
                                                Langage Java /
                                                Gestionnaire de dépendances Maven / 
                                                Gestionnaire de dépendances Gradle /
                                                Librairie Spring /
                                                Librairie Jsoup /
                                                Librairie Selenium WebDriver
                                            </li>
                                            <li>
                                                Langage JavaScript /
                                                Langage Visual Basic /
                                                Langage C# /
                                                Langage TypeScript
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Back-end -->

                        <br>

                        <!-- Tests unitaires et qualité -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Tests unitaires et qualité
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Package Unittests avec Python / 
                                                Package JUnit avec Java / 
                                                Mockito avec Java / 
                                                WireMock avec Java
                                            </li>
                                            <li>
                                                Logiciel SonarQube
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Tests unitaires et qualité -->

                        <br>

                        <!-- Systèmes de bases de données -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Systèmes de bases de données
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Serveur de bases de données MySQL / 
                                                Logiciel MySQL Workbench
                                            </li>
                                            <li>
                                                Serveur de bases de données graphe de Neo4j / 
                                                Serveur de bases de données documents MongoDB
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Systèmes de bases de données -->

                        <br>

                        <!-- Outils de déploiement -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Outils de déploiement
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Docker / Git
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Outils de déploiement -->

                        <br>

                        <!-- Intelligence artificielle -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Intelligence artificielle
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Exploration de données / 
                                                Gestion de la connaissance / 
                                                Automatisation de processus robotisés /
                                                RPA
                                            </li>
                                            <li>
                                                Cloud computing [Open data; Autodesk Forge; Netheos] / 
                                                Reconnaissance vocale / 
                                                Vision par ordinateur
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Intelligence artificielle -->

                        <br>

                        <!-- Mobile -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Mobile
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Système d'exploitation Android / 
                                                Logiciel Android Studio
                                            </li>
                                            <li>
                                                Framework Kivy pour le développement mobile multiplateforme développé 
                                                en Python
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Mobile -->

                        <br>

                        <!-- Systèmes embarqués -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Systèmes embarqués
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Carte embarquée Raspberry Pi / 
                                                Système d’exploitation Raspbian
                                            </li>
                                            <li>
                                                Cartes embarquées Arduino Uno et Arduino Yun /
                                                Système d’exploitation OpenWRT
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Systèmes embarqués -->

                        <br>

                        <!-- Conception assistée par ordinateur -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Conception assistée par ordinateur
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Logiciel FreeCAD développé en Python pour la conception de pièces 
                                                mécaniques de précision (https://www.freecadweb.org/?lang=fr)
                                            </li>
                                            <li>
                                                Logiciel Skidl développé en Python pour la génération de fichiers 
                                                netlist pour la fabrication de circuits imprimés
                                                (https://devbisme.github.io/skidl/)
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Conception assistée par ordinateur -->

                        <br>

                        <!-- Gestion de version -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Gestion de version
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Logiciel de gestion de version Git / 
                                                Gestion de mon compte GitHub : https://github.com/Jay4C
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Gestion de version -->

                        <br>

                        <!-- Réseaux et Systèmes -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Réseaux et systèmes
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Système d’exploitation Windows / 
                                                Système d’exploitation Linux
                                            </li>
                                            <li>
                                                Protocoles informatiques [HTTP, HTTPS, TCP, IPv4, IPv6, SSH, FTP, 
                                                SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, Modbus RTU]
                                            </li>
                                            <li>
                                                Librairie MinimalModbus avec Python / 
                                                Librairie Jlibmodbus avec Java
                                            </li>
                                            <li>
                                                Routeur 4G / 
                                                Port Forwarding / 
                                                DynDNS / 
                                                Dongle 3G
                                            </li>
                                            <li>
                                                Logiciel Wireshark / 
                                                Logiciel PuTTY / 
                                                Logiciel WinSCP /
                                                Logiciel Postman /
                                                Logiciel TeamViewer / 
                                                Logiciel VirtualBox /
                                                Logiciel Tor /
                                                Logiciel Angry IP Scanner
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Réseaux et Systèmes -->

                        <br>

                        <!-- Gestion de projet et communications -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Gestion de projet et communications
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Brainstorming / 
                                                Logiciel FreeMind / 
                                                Logiciel Pencil / 
                                                Diagrammes UML / 
                                                Logiciel ArgoUML
                                            </li>
                                            <li>
                                                Logiciel GanttProject / 
                                                Planification des étapes d'un projet
                                            </li>
                                            <li>
                                                Management de projet / 
                                                Management des opérations / 
                                                Management Agile / 
                                                Management de la qualité
                                            </li>
                                            <li>
                                                Analyse des besoins du client / 
                                                Établissement d’un cahier des charges
                                            </li>
                                            <li>
                                                Rédaction d’un support technique / 
                                                Rédaction d'une base de données de connaissances
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Gestion de projet et communications -->

                        <br>

                        <!-- Droit -->
                        <table class="table table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">
                                        Droit
                                    </th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Code général des impôts / 
                                                Code monétaire et financier / 
                                                Code de l'environnement / 
                                                Code de l'énergie
                                            </li>
                                            <li>
                                                Code des relations entre le public et l’administration / 
                                                Code de la construction et de l'habitation
                                            </li>
                                            <li>
                                                Code de l'urbanisme /
                                                Code des postes et des communications électroniques /                                                 
                                            </li>
                                            <li>
                                                Site internet "www.legifrance.gouv.fr" /
                                                Site internet "www.lexadin.nl" /
                                                Site internet "https://europa.eu/european-union/law/find-legislation_fr"
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Droit -->
                    </div>
                </div>
                <!-- Compétences -->

                <br>

                <!-- Expériences -->
                <div class="card">
                    <div class="card-header">
                        <h4>
                            Expériences
                        </h4>
                    </div>
                    <div class="card-body">
                        <!-- Holomorphe -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Développement informatique polyvalent
                                        <br>
                                        Période : Mai 2020 – En cours
                                        <br>
                                        Statut professionnel : Président
                                        <br>
                                        Société : Holomorphe - Paris
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement d'applications Python pour créer des pièces mécaniques 
                                                de précision pour les technologies de l'hydrogène avec le logiciel
                                                FreeCAD disponibles sur https://github.com/Jay4C/Python-Macros-For-FreeCAD
                                            </li>
                                            <li>
                                                Développement d'applications Python pour l'automatisation de processus 
                                                robotisés disponibles sur https://github.com/Jay4C/Web-Automation
                                            </li>
                                            <li>
                                                Développement d'une application Python pour récupérer toutes 
                                                les parcelles convenables pour injecter du méthane de synthèse dans un 
                                                réseau de distribution de gaz naturel en France
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Holomorphe -->

                        <br>

                        <!-- Société PASS Technologie -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Développeur informatique polyvalent
                                        <br>
                                        Période : Février 2019 - Avril 2019 [3 mois]
                                        <br>
                                        Statut professionnel : Salarié
                                        <br>
                                        Société : PASS Technologie - Paris
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement d'une application pour le BIM (Building Information
                                                Modeling)
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Société Pass Tech -->

                        <br>

                        <!-- Kriir -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Ingénieur d'études et développement
                                        <br>
                                        Période : Septembre – Novembre 2018 [3 mois]
                                        <br>
                                        Statut professionnel : Salarié
                                        <br>
                                        Société : Kriir - Ermont
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement d’une application web en Java pour utiliser la signature
                                                électronique
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Kriir -->

                        <br>

                        <!-- Refuel -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Lecture et écriture sur un automate en modbus
                                        <br>
                                        Période : Juin - Juillet 2018 [2 mois]
                                        <br>
                                        Statut professionnel : Indépendant
                                        <br>
                                        Société : Refuel S.A.S. - Trappes
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement de plusieurs fonctionnalités permettant de lire et
                                                d’écrire sur un automate programmable via le
                                                protocole Modbus TCP avec le langage Java en liaison WiFi
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Refuel  -->

                        <br>

                        <!-- SETEC SMART EFFICIENCY -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Génération automatique de rapport énergétique
                                        <br>
                                        Période : Novembre - Décembre 2017 [2 mois]
                                        <br>
                                        Statut professionnel : Indépendant
                                        <br>
                                        Société : Setec Smart Efficiency - Paris
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développer un programme Java pour générer un rapport énergétique
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- SETEC SMART EFFICIENCY -->

                        <br>

                        <!--  SOCIÉTÉ GAMELOFT -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Vérification de la pertinence des données
                                        <br>
                                        Période : Juin - Juillet 2017 [2 mois]
                                        <br>
                                        Statut professionnel : Indépendant
                                        <br>
                                        Société : Gameloft - Paris
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Développement d’une macro Visual Basic pour un fichier Excel 
                                                permettant de contrôler la cohérence des données saisies dans 
                                                plusieurs fichiers Excel et consolider ces données dans un seul 
                                                fichier Excel pour les ressources humaines
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- SOCIÉTÉ GAMELOFT -->

                        <br>
                    </div>
                </div>
                <!-- Expériences -->

                <br>

                <!-- Formations -->
                <div class="card">
                    <div class="card-header">
                        <h4>
                            Formations
                        </h4>
                    </div>
                    <div class="card-body">
                        <!-- Formation Java -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Formation Java / JEE
                                        <br>
                                        Période : Avril - Juillet 2017 [4 mois]
                                        <br>
                                        Centre de formation : INTI Formation – Paris Tour Montparnasse
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Apprentissage du langage informatique Java / Apprentissage de la
                                                conception orientée objet
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Formation Java -->

                        <br>

                        <!-- Diplome ingenieur -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <td scope="col">
                                        Intitulé : Diplôme d'ingénieur généraliste
                                        <br>
                                        Période : Septembre 2013 - Août 2016 [3 ans]
                                        <br>
                                        Centre de formation : École Supérieure d'Ingénieurs Léonard de Vinci – Paris La
                                        Défense
                                    </td>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        <ul>
                                            <li>
                                                Nouvelles énergies / Informatique / Finance / Mécanique numérique
                                            </li>
                                        </ul>
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Diplome ingenieur -->
                    </div>
                </div>
                <!-- Formations -->

                <br>

                <!-- Centre d'intérêts -->
                <div class="card">
                    <div class="card-header">
                        <h4>
                            Centre d'intérêts
                        </h4>
                    </div>
                    <div class="card-body">
                        <ul>
                            <li>
                                Veille technologique / 
                                Brevets d'inventions
                            </li>
                            <li>
                                Droit français / 
                                Droit européen / 
                                Droit international
                            </li>
                            <li>
                                Physique / 
                                Mathématiques / 
                                Finance / 
                                Economie / 
                                Management / 
                                Marketing / 
                                Spiritualité / 
                                Ufologie
                            </li>
                        </ul>
                    </div>
                </div>
                <!-- Centre d'intérêts -->
            </div>
        </div>
    </div>

    <br>

    <!-- Optional JavaScript -->
    <!-- jQuery first, then Popper.js, then Bootstrap JS -->
    <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
        integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
        crossorigin="anonymous"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
        integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
        crossorigin="anonymous"></script>
    <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
        integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
        crossorigin="anonymous"></script>
</body>

</html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de Monsieur ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Employé polyvalent / Développeur informatique'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body, "CV\\CV_Employe_Polyvalent_[Mr_].pdf", configuration=config,
                           options=options)

    # ok
    def test_cv_dev_python__ingenieur_energie_renouvelable(self):
        print("test_cv_ingenieur_energie_renouvelable")

        body = """
        <!doctype html>
        <html lang="en">

        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

            <title>CV de  - Développeur Python / Ingénieur en énergies renouvelables</title>
        </head>

        <body>
            <br>
            <div class="container">
                <div class="card text-center">
                    <div class="card-header text-center">
                        <h2>
                            Développeur Python / Ingénieur en énergies renouvelables
                        </h2>
                    </div>
                    <div class="card-body text-center">

                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">Identité</th>
                                    <th scope="col">Adresse postale</th>
                                    <th scope="col">Téléphone</th>
                                    <th scope="col">Email</th>
                                    <th scope="col">Nationalité</th>
                                    <th scope="col">Age</th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        
                                    </td>
                                    <td>
                                         93800 Epinay-sur-Seine (France)
                                    </td>
                                    <td>
                                        
                                    </td>
                                    <td>
                                        @outlook.fr
                                    </td>
                                    <td>
                                        Français
                                    </td>
                                    <td>
                                        28 ans
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Informations personnelles -->

                        <br>

                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Compétences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Energie -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Energie
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- Type of machines -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Type d'appareils
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Panneaux photovoltaïques / Éolienne / Turbine à gaz 
                                                                / Fusion nucléaire / Fusion froide / 
                                                                Machines à énergie libre
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Type of machines -->

                                        <br>

                                        <!-- Mathematics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mathématique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Nombres complexes / Géométrie élémentaire du plan / 
                                                                Géométrie élémentaire de l'espace / Fonctions usuelles 
                                                                / Équations différentielles linéaires / 
                                                                Étude des courbes planes / Conique / 
                                                                Entiers naturels, ensembles finis, nombres 
                                                                / Corps de nombres réels / Suites de nombres réels 
                                                                / Fonctions d'un réel variable réelle à valeurs 
                                                                multiples / Dérivation de fonctions à valeurs réelles 
                                                                / Intégration de fonctions à valeurs réelles sur 
                                                                un segment / Développements limités / Propriétés 
                                                                métriques des arcs / Suites et fonctions à valeurs 
                                                                complexes / Notions sur les fonctions de deux 
                                                                variables réelles / Intégrales multiples / Algébrique 
                                                                structures / Arithmétique / Polynômes / Fractions 
                                                                rationnelles / Espaces vectoriels / Dimension des 
                                                                espaces vectoriels / Calcul matriciel / Groupe 
                                                                symétrique, déterminant / Produit scalaire, 
                                                                groupe orthogonal / Géométrie affine / 
                                                                Techniques de démonstration / Techniques d'algèbre / 
                                                                Techniques d'analyse / Analyse multivariable / 
                                                                Théorie des graphes / Théorie de l'information / 
                                                                Analyse des données
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mathematics -->

                                        <br>

                                        <!-- Physics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Physique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Thermodynamique / Electrotechnique / Electromagnétisme 
                                                                / Electronique / Thermo-électricité / Mécanique 
                                                                quantique / Physique nucléaire / Electricité / 
                                                                Physique des plasmas / Résonance magnétique nucléaire / 
                                                                Matériaux diélectriques / Mécanique des fluides / 
                                                                Oscillateurs harmoniques et signaux sinusoïdaux / 
                                                                Circuits linéaires en régime continu / Dispositions 
                                                                transitoires du premier ordre / Régime de transition 
                                                                du second ordre / Régime sinusoïdal forcé - résonance / 
                                                                Filtrage / Cinématique / Principe fondamental de la 
                                                                dynamique / Énergie mécanique / Mouvement dans un champ 
                                                                électrique et magnétique / Propagation du signal - 
                                                                notion d'ondes / Lois de Snell-Descartes - réflexion 
                                                                et réfraction / Lentilles minces sphériques / Moment 
                                                                cinétique - solide tournant autour d'un axe fixe / 
                                                                Forces centrales conservatrices / États de la matière 
                                                                / Machines thermiques / Champ magnétique - forces de 
                                                                Laplace - induction
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Physics -->

                                        <br>

                                        <!-- Management -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Management
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Management collaboratif / Management commercial / 
                                                                Management conseil / Management d'équipe / 
                                                                Management de l'innovation / Management de projet / 
                                                                Management de la qualité / Management des connaissances 
                                                                / Management des opérations / Management des ressources 
                                                                humaines / Management des risques / Management du 
                                                                changement / Management évènementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Management -->

                                        <br>

                                        <!-- Marketing -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Marketing
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Marketing commercial / Marketing B2B / Marketing 
                                                                digital / Marketing direct / Marketing opérationnel / 
                                                                Marketing relationnel / Marketing stratégique / 
                                                                Marketing événementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Marketing -->

                                        <br>

                                        <!-- Finance -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Finance
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Analyse financière / Finance d'entreprise / Finance 
                                                                internationale / Finance personnelle / Financement de 
                                                                projet / Forex
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Finance -->

                                        <br>

                                        <!-- Economy -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Economie
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Macroéconomie / Microéconomie / Agents économiques 
                                                                / Économie agricole / Économie circulaire / 
                                                                Consommation et épargne / Économie de l'éducation / 
                                                                Économie financière / Économie géographique / 
                                                                Économie industrielle / Économie globale / 
                                                                Économie monétaire
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Economy -->

                                        <br>

                                        <!-- International law -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Loi internationale
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Site internet : "www.legifrance.gouv.fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "www.lexadin.nl"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://europa.eu/european-union/law/find-legislation_fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://www.fedlex.admin.ch/fr/home"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://legislation.gov.ky/"
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- International law -->        
                                    </div>
                                </div>
                                <!-- Energy -->

                                <br>

                                <!-- Computing -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Informatique
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- preferred programming languages -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Langages de programmation préférés
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python / Java / JavaScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Preferred programming languages -->    

                                        <br>

                                        <!-- Office -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Bureautique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Microsoft Excel, Word, PowerPoint, Access, 
                                                                Outlook
                                                            </li>
                                                            <li>
                                                                Open Office Calc, Writer
                                                            </li>
                                                            <li>
                                                                Intellij pour Java, PyCharm pour Python
                                                            </li>
                                                            <li>
                                                                Logiciel Pencil pour faire des schémas simplifiés / 
                                                                Logiciel GanttProject pour la gestion de projet
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Office -->

                                        <br>

                                        <!-- Front-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Front-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                HTML markup language /
                                                                CSS /
                                                                JQuery Framework /
                                                                Django Framework /
                                                                MJML Framework
                                                            </li>
                                                            <li>
                                                                Visual Studio Software Code /
                                                                Angular Framework 7 /
                                                                Boostrap /
                                                                Leaflet.js
                                                            </li>
                                                            <li>
                                                                Vis.js / 
                                                                Artyom.js /
                                                                Flowchart.js
                                                            </li>
                                                            <li>
                                                                Tomcat server /
                                                                Apache server /
                                                                Nginx server
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Front-end -->

                                        <br>

                                        <!-- Back-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Back-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python /
                                                                Flask /
                                                                BeautifulSoup /
                                                                pdfkit /
                                                                xlswriter /
                                                                xlrd /
                                                                pywinauto
                                                            </li>
                                                            <li>
                                                                Java /
                                                                Maven / 
                                                                Gradle /
                                                                Spring /
                                                                Jsoup /
                                                                Selenium WebDriver
                                                            </li>
                                                            <li>
                                                                JavaScript /
                                                                Visual Basic /
                                                                C# /
                                                                TypeScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Back-end -->

                                        <br>

                                        <!-- Tests unitaires et qualité -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Tests unitaires et qualité
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Unittests with Python / 
                                                                JUnit with Java / 
                                                                Mockito with Java / 
                                                                WireMock with Java
                                                            </li>
                                                            <li>
                                                                SonarQube
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Tests unitaires et qualité -->

                                        <br>

                                        <!-- Systèmes de bases de données -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes de base de données
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                MySQL / 
                                                                MySQL Workbench
                                                            </li>
                                                            <li>
                                                                NeO4J, serveur de base de données orienté graphe  / 
                                                                MongoDB, serveur de base de données orienté document
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes de bases de données -->

                                        <br>

                                        <!-- Outils de déploiement -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        DevOps
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Docker / Git
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Outils de déploiement -->

                                        <br>

                                        <!-- Intelligence artificielle -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Intelligence artificielle
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Exploration de données / 
                                                                Gestion des connaissances / 
                                                                Automatisation des processus robotiques /
                                                                Cloud computing / Reconnaissance vocale / 
                                                                Vision par ordinateur
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Intelligence artificielle -->

                                        <br>

                                        <!-- Mobile -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mobile
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Android /
                                                                Kivy
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mobile -->

                                        <br>

                                        <!-- Systèmes embarqués -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes embarqués
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Raspberry Pi / 
                                                                Arduino Uno and Arduino Yun / 
                                                                OpenWRT OS
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes embarqués -->

                                        <br>

                                        <!-- Conception assistée par ordinateur -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Conception assistée par ordinateur
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                FreeCAD / 
                                                                Skidl /
                                                                KiCAD
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Conception assistée par ordinateur -->

                                        <br>

                                        <!-- Gestion de version -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion des versions
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Git / 
                                                                GitHub : https://github.com/Jay4C
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de version -->

                                        <br>

                                        <!-- Réseaux et Systèmes -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Réseaux et systèmes
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Windows / 
                                                                Linux
                                                            </li>
                                                            <li>
                                                                Protocoles informatiques [HTTP, HTTPS, TCP, IPv4, IPv6, 
                                                                SSH, FTP, SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, 
                                                                Modbus RTU]
                                                            </li>
                                                            <li>
                                                                MinimalModbus avec Python / 
                                                                Jlibmodbus avec Java
                                                            </li>
                                                            <li>
                                                                Routeur 4G / 
                                                                Port Forwarding / 
                                                                DynDNS / 
                                                                Dongle 3G
                                                            </li>
                                                            <li>
                                                                Wireshark / 
                                                                PuTTY / 
                                                                WinSCP /
                                                                Postman /
                                                                TeamViewer / 
                                                                VirtualBox /
                                                                Tor /
                                                                Angry IP Scanner
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Réseaux et Systèmes -->

                                        <br>

                                        <!-- Gestion de projet et communications -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion de projet et communication
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Brainstorming / 
                                                                FreeMind / 
                                                                Pencil / 
                                                                Diagrammes UML/ 
                                                                ArgoUML
                                                            </li>
                                                            <li>
                                                                GanttProject / 
                                                                Planification des étapes d'un projet
                                                            </li>
                                                            <li>
                                                                Management de projet / Management des opérations 
                                                                / Management agile / Management de la qualité
                                                            </li>
                                                            <li>
                                                                Analyse des besoins client / 
                                                                Etablissement d'un cahier des charges
                                                            </li>
                                                            <li>
                                                                Rédaction d'un support technique / 
                                                                Rédaction d'une base de connaissances
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de projet et communications -->        
                                    </div>
                                </div>
                                <!-- Computing -->
                            </div>
                        </div>
                        <!-- Compétences -->

                        <br>

                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Société PASS Technologie -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Développeur informatique polyvalent
                                                <br>
                                                Période : février 2019 - avril 2019 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Société : Pass Technology - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application pour le BIM (Building 
                                                        Information Modeling)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Pass Tech -->

                                <br>

                                <!-- Kriir -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Ingénieur informatique
                                                <br>
                                                Période : septembre - novembre 2018 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Entreprise : Kriir - Ermont
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application web Java pour utiliser la 
                                                        signature électronique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Kriir -->

                                <br>

                                <!-- Refuel -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Lire et écrire sur un automate en Modbus
                                                <br>
                                                Période : juin - juillet 2018 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Refuel S.A.S. - Trappes
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de plusieurs fonctionnalités pour lire et écrire 
                                                        sur un automate programmable via le protocole Modbus TCP avec 
                                                        langage Java en liaison WiFi
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Refuel  -->

                                <br>

                                <!-- SETEC SMART EFFICIENCY -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Génération automatique de rapport énergétique
                                                <br>
                                                Période : novembre - décembre 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : SETEC SMART EFFICIENCY - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développer un programme en Java pour générer des rapports 
                                                        énergétique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SETEC SMART EFFICIENCY -->

                                <br>

                                <!--  SOCIÉTÉ GAMELOFT -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Vérification de la pertinence des données
                                                <br>
                                                Période : juin - juillet 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Gameloft - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une macro Visual Basic pour un fichier 
                                                        Excel permettant de contrôler la cohérence des données saisies 
                                                        dans plusieurs fichiers Excel et de consolider ces données 
                                                        dans un seul fichier Excel pour les ressources humaines
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SOCIÉTÉ GAMELOFT -->
                            </div>
                        </div>
                        <!-- Expériences -->

                        <br>

                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Formation Java -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Formation Java / JEE
                                                <br>
                                                Période : avril - juillet 2017 [4 mois]
                                                <br>
                                                Centre de formation : INTI Formation - Paris Tour Montparnasse
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Apprentissage du langage informatique Java / 
                                                        Apprentissage de la programmation orientée objet
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Formation Java -->

                                <br>

                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : septembre 2013 - août 2016 [3 ans]
                                                <br>
                                                Centre de formation : Ecole Supérieure d'Ingénieurs Léonard de Vinci 
                                                - Paris La Défense (France)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / Informatique / Finance / Mécanique 
                                                        numérique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->

                        <br>

                        <!-- Languages -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Langues
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Anglais écrit et oral [TOEIC : 840 points ; Date d'expiration : 11 juillet 2020]
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Languages -->

                        <br>

                        <!-- Centre d'intérêts -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Centre d'intérêts
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Technologie / Brevets
                                    </li>
                                    <li>
                                        Droit français / 
                                        Droit européen / 
                                        Droit international
                                    </li>
                                    <li>
                                        Physique / 
                                        Mathématiques / 
                                        Finance / 
                                        Economie / 
                                        Gestion / 
                                        Marketing / 
                                        Spiritualité / 
                                        Ufologie
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Centre d'intérêts -->
                    </div>
                </div>
            </div>

            <br>

            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>
        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Développeur Python / Ingénieur en énergies renouvelables'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body, "CV\\CV__[Dev_Python__Ingenieur_Energies_Renouvelables].pdf", configuration=config,
                           options=options)


class CV_with_Docx(unittest.TestCase):
    #
    def test_cv_developpeur_python(self):
        print("test_cv_developpeur_python")

        from docx import Document
        from docx.shared import Inches

        document = Document()

        document.add_heading('Document Title', 0)

        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True


class CV_Ingenieur_Energies_Renouvelables(unittest.TestCase):
    # ok
    def test_cv_ingenieur_energie_renouvelable(self):
        print("test_cv_ingenieur_energie_renouvelable")

        body = """
        <!doctype html>
        <html lang="en">

        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

            <title>CV de  - Ingénieur en énergies renouvelables</title>
        </head>

        <body>
            <br>
            <div class="container">
                <div class="card text-center">
                    <div class="card-header text-center">
                        <h2>
                            Ingénieur en énergies renouvelables
                        </h2>
                    </div>
                    <div class="card-body text-center">

                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">Identité</th>
                                    <th scope="col">Adresse postale</th>
                                    <th scope="col">Téléphone</th>
                                    <th scope="col">Email</th>
                                    <th scope="col">Age</th>
                                    <th scope="col">Nationalité</th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        
                                    </td>
                                    <td>
                                         93800 Epinay-sur-Seine (France)
                                    </td>
                                    <td>
                                        
                                    </td>
                                    <td>
                                        @outlook.fr
                                    </td>
                                    <td>
                                        28 ans
                                    </td>
                                    <td>
                                        Français
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Informations personnelles -->

                        <br>

                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Compétences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Energie -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Energie
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- Type of machines -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Type d'appareils
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Panneaux photovoltaïques / Éolienne / Turbine à gaz 
                                                                / Fusion nucléaire / Fusion froide / 
                                                                Machines à énergie libre
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Type of machines -->

                                        <br>

                                        <!-- Mathematics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mathématique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Nombres complexes / Géométrie élémentaire du plan / 
                                                                Géométrie élémentaire de l'espace / Fonctions usuelles 
                                                                / Équations différentielles linéaires / 
                                                                Étude des courbes planes / Conique / 
                                                                Entiers naturels, ensembles finis, nombres 
                                                                / Corps de nombres réels / Suites de nombres réels 
                                                                / Fonctions d'un réel variable réelle à valeurs 
                                                                multiples / Dérivation de fonctions à valeurs réelles 
                                                                / Intégration de fonctions à valeurs réelles sur 
                                                                un segment / Développements limités / Propriétés 
                                                                métriques des arcs / Suites et fonctions à valeurs 
                                                                complexes / Notions sur les fonctions de deux 
                                                                variables réelles / Intégrales multiples / Algébrique 
                                                                structures / Arithmétique / Polynômes / Fractions 
                                                                rationnelles / Espaces vectoriels / Dimension des 
                                                                espaces vectoriels / Calcul matriciel / Groupe 
                                                                symétrique, déterminant / Produit scalaire, 
                                                                groupe orthogonal / Géométrie affine / 
                                                                Techniques de démonstration / Techniques d'algèbre / 
                                                                Techniques d'analyse / Analyse multivariable / 
                                                                Théorie des graphes / Théorie de l'information / 
                                                                Analyse des données
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mathematics -->

                                        <br>

                                        <!-- Physics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Physique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Thermodynamique / Electrotechnique / Electromagnétisme 
                                                                / Electronique / Thermo-électricité / Mécanique 
                                                                quantique / Physique nucléaire / Electricité / 
                                                                Physique des plasmas / Résonance magnétique nucléaire / 
                                                                Matériaux diélectriques / Mécanique des fluides / 
                                                                Oscillateurs harmoniques et signaux sinusoïdaux / 
                                                                Circuits linéaires en régime continu / Dispositions 
                                                                transitoires du premier ordre / Régime de transition 
                                                                du second ordre / Régime sinusoïdal forcé - résonance / 
                                                                Filtrage / Cinématique / Principe fondamental de la 
                                                                dynamique / Énergie mécanique / Mouvement dans un champ 
                                                                électrique et magnétique / Propagation du signal - 
                                                                notion d'ondes / Lois de Snell-Descartes - réflexion 
                                                                et réfraction / Lentilles minces sphériques / Moment 
                                                                cinétique - solide tournant autour d'un axe fixe / 
                                                                Forces centrales conservatrices / États de la matière 
                                                                / Machines thermiques / Champ magnétique - forces de 
                                                                Laplace - induction
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Physics -->

                                        <br>

                                        <!-- Management -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Management
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Management collaboratif / Management commercial / 
                                                                Management conseil / Management d'équipe / 
                                                                Management de l'innovation / Management de projet / 
                                                                Management de la qualité / Management des connaissances 
                                                                / Management des opérations / Management des ressources 
                                                                humaines / Management des risques / Management du 
                                                                changement / Management évènementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Management -->

                                        <br>

                                        <!-- Marketing -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Marketing
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Marketing commercial / Marketing B2B / Marketing 
                                                                digital / Marketing direct / Marketing opérationnel / 
                                                                Marketing relationnel / Marketing stratégique / 
                                                                Marketing événementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Marketing -->

                                        <br>

                                        <!-- Finance -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Finance
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Analyse financière / Finance d'entreprise / Finance 
                                                                internationale / Finance personnelle / Financement de 
                                                                projet / Forex
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Finance -->

                                        <br>

                                        <!-- Economy -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Economie
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Macroéconomie / Microéconomie / Agents économiques 
                                                                / Économie agricole / Économie circulaire / 
                                                                Consommation et épargne / Économie de l'éducation / 
                                                                Économie financière / Économie géographique / 
                                                                Économie industrielle / Économie globale / 
                                                                Économie monétaire
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Economy -->

                                        <br>

                                        <!-- International law -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Loi internationale
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Site internet : "www.legifrance.gouv.fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "www.lexadin.nl"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://europa.eu/european-union/law/find-legislation_fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://www.fedlex.admin.ch/fr/home"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://legislation.gov.ky/"
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- International law -->        
                                    </div>
                                </div>
                                <!-- Energy -->

                                <br>

                                <!-- Computing -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Informatique
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- preferred programming languages -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Langages de programmation préférés
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python / Java / JavaScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Preferred programming languages -->    

                                        <br>

                                        <!-- Office -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Bureautique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Microsoft Excel, Word, PowerPoint, Access, 
                                                                Outlook
                                                            </li>
                                                            <li>
                                                                Open Office Calc, Writer
                                                            </li>
                                                            <li>
                                                                Intellij pour Java, PyCharm pour Python
                                                            </li>
                                                            <li>
                                                                Logiciel Pencil pour faire des schémas simplifiés / 
                                                                Logiciel GanttProject pour la gestion de projet
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Office -->

                                        <br>

                                        <!-- Front-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Front-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                HTML markup language /
                                                                CSS /
                                                                JQuery Framework /
                                                                Django Framework /
                                                                MJML Framework
                                                            </li>
                                                            <li>
                                                                Visual Studio Software Code /
                                                                Angular Framework 7 /
                                                                Boostrap /
                                                                Leaflet.js
                                                            </li>
                                                            <li>
                                                                Vis.js / 
                                                                Artyom.js /
                                                                Flowchart.js
                                                            </li>
                                                            <li>
                                                                Tomcat server /
                                                                Apache server /
                                                                Nginx server
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Front-end -->

                                        <br>

                                        <!-- Back-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Back-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python /
                                                                Flask /
                                                                BeautifulSoup /
                                                                pdfkit /
                                                                xlswriter /
                                                                xlrd /
                                                                pywinauto
                                                            </li>
                                                            <li>
                                                                Java /
                                                                Maven / 
                                                                Gradle /
                                                                Spring /
                                                                Jsoup /
                                                                Selenium WebDriver
                                                            </li>
                                                            <li>
                                                                JavaScript /
                                                                Visual Basic /
                                                                C# /
                                                                TypeScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Back-end -->

                                        <br>

                                        <!-- Tests unitaires et qualité -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Tests unitaires et qualité
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Unittests with Python / 
                                                                JUnit with Java / 
                                                                Mockito with Java / 
                                                                WireMock with Java
                                                            </li>
                                                            <li>
                                                                SonarQube
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Tests unitaires et qualité -->

                                        <br>

                                        <!-- Systèmes de bases de données -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes de base de données
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                MySQL / 
                                                                MySQL Workbench
                                                            </li>
                                                            <li>
                                                                NeO4J, serveur de base de données orienté graphe  / 
                                                                MongoDB, serveur de base de données orienté document
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes de bases de données -->

                                        <br>

                                        <!-- Outils de déploiement -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        DevOps
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Docker / Git
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Outils de déploiement -->

                                        <br>

                                        <!-- Intelligence artificielle -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Intelligence artificielle
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Exploration de données / 
                                                                Gestion des connaissances / 
                                                                Automatisation des processus robotiques /
                                                                Cloud computing / Reconnaissance vocale / 
                                                                Vision par ordinateur
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Intelligence artificielle -->

                                        <br>

                                        <!-- Mobile -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mobile
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Android /
                                                                Kivy
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mobile -->

                                        <br>

                                        <!-- Systèmes embarqués -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes embarqués
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Raspberry Pi / 
                                                                Arduino Uno and Arduino Yun / 
                                                                OpenWRT OS
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes embarqués -->

                                        <br>

                                        <!-- Conception assistée par ordinateur -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Conception assistée par ordinateur
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                FreeCAD / 
                                                                Skidl /
                                                                KiCAD
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Conception assistée par ordinateur -->

                                        <br>

                                        <!-- Gestion de version -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion des versions
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Git / 
                                                                GitHub : https://github.com/Jay4C
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de version -->

                                        <br>

                                        <!-- Réseaux et Systèmes -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Réseaux et systèmes
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Windows / 
                                                                Linux
                                                            </li>
                                                            <li>
                                                                Protocoles informatiques [HTTP, HTTPS, TCP, IPv4, IPv6, 
                                                                SSH, FTP, SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, 
                                                                Modbus RTU]
                                                            </li>
                                                            <li>
                                                                MinimalModbus avec Python / 
                                                                Jlibmodbus avec Java
                                                            </li>
                                                            <li>
                                                                Routeur 4G / 
                                                                Port Forwarding / 
                                                                DynDNS / 
                                                                Dongle 3G
                                                            </li>
                                                            <li>
                                                                Wireshark / 
                                                                PuTTY / 
                                                                WinSCP /
                                                                Postman /
                                                                TeamViewer / 
                                                                VirtualBox /
                                                                Tor /
                                                                Angry IP Scanner
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Réseaux et Systèmes -->

                                        <br>

                                        <!-- Gestion de projet et communications -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion de projet et communication
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Brainstorming / 
                                                                FreeMind / 
                                                                Pencil / 
                                                                Diagrammes UML/ 
                                                                ArgoUML
                                                            </li>
                                                            <li>
                                                                GanttProject / 
                                                                Planification des étapes d'un projet
                                                            </li>
                                                            <li>
                                                                Management de projet / Management des opérations 
                                                                / Management agile / Management de la qualité
                                                            </li>
                                                            <li>
                                                                Analyse des besoins client / 
                                                                Etablissement d'un cahier des charges
                                                            </li>
                                                            <li>
                                                                Rédaction d'un support technique / 
                                                                Rédaction d'une base de connaissances
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de projet et communications -->        
                                    </div>
                                </div>
                                <!-- Computing -->
                            </div>
                        </div>
                        <!-- Compétences -->

                        <br>

                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Société PASS Technologie -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Développeur informatique polyvalent
                                                <br>
                                                Période : février 2019 - avril 2019 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Société : Pass Technology - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application pour le BIM (Building 
                                                        Information Modeling)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Pass Tech -->

                                <br>

                                <!-- Kriir -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Ingénieur informatique
                                                <br>
                                                Période : septembre - novembre 2018 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Entreprise : Kriir - Ermont
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application web Java pour utiliser la 
                                                        signature électronique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Kriir -->

                                <br>

                                <!-- Refuel -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Lire et écrire sur un automate en Modbus
                                                <br>
                                                Période : juin - juillet 2018 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Refuel S.A.S. - Trappes
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de plusieurs fonctionnalités pour lire et écrire 
                                                        sur un automate programmable via le protocole Modbus TCP avec 
                                                        langage Java en liaison WiFi
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Refuel  -->

                                <br>

                                <!-- SETEC SMART EFFICIENCY -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Génération automatique de rapport énergétique
                                                <br>
                                                Période : novembre - décembre 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : SETEC SMART EFFICIENCY - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développer un programme en Java pour générer des rapports 
                                                        énergétique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SETEC SMART EFFICIENCY -->

                                <br>

                                <!--  SOCIÉTÉ GAMELOFT -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Vérification de la pertinence des données
                                                <br>
                                                Période : juin - juillet 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Gameloft - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une macro Visual Basic pour un fichier 
                                                        Excel permettant de contrôler la cohérence des données saisies 
                                                        dans plusieurs fichiers Excel et de consolider ces données 
                                                        dans un seul fichier Excel pour les ressources humaines
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SOCIÉTÉ GAMELOFT -->
                            </div>
                        </div>
                        <!-- Expériences -->

                        <br>

                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Formation Java -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Formation Java / JEE
                                                <br>
                                                Période : avril - juillet 2017 [4 mois]
                                                <br>
                                                Centre de formation : INTI Formation - Paris Tour Montparnasse
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Apprentissage du langage informatique Java / 
                                                        Apprentissage de la programmation orientée objet
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Formation Java -->

                                <br>

                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : septembre 2013 - août 2016 [3 ans]
                                                <br>
                                                Centre de formation : Ecole Supérieure d'Ingénieurs Léonard de Vinci 
                                                - Paris La Défense (France)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / Informatique / Finance / Mécanique 
                                                        numérique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->

                        <br>

                        <!-- Languages -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Langues
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Anglais écrit et oral [TOEIC : 840 points ; Date d'expiration : 11 juillet 2020]
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Languages -->

                        <br>

                        <!-- Centre d'intérêts -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Centre d'intérêts
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Technologie / Brevets
                                    </li>
                                    <li>
                                        Droit français / 
                                        Droit européen / 
                                        Droit international
                                    </li>
                                    <li>
                                        Physique / 
                                        Mathématiques / 
                                        Finance / 
                                        Economie / 
                                        Gestion / 
                                        Marketing / 
                                        Spiritualité / 
                                        Ufologie
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Centre d'intérêts -->
                    </div>
                </div>
            </div>

            <br>

            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>
        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Ingénieur en énergies renouvelables'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body, "CV\\CV__[Ingenieur_Energies_Renouvelables].pdf", configuration=config,
                           options=options)

    # ok
    def test_cv_ingenieur_energie_renouvelable_v1(self):
        print("test_cv_ingenieur_energie_renouvelable_v1")

        body = """
        <!doctype html>
        <html lang="en">

        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

            <title>CV de  - Ingénieur en énergies renouvelables</title>
        </head>

        <body>
            <br>
            <div class="container">
                <div class="card text-center">
                    <div class="card-header text-center">
                        <h2>
                            Ingénieur en énergies renouvelables
                        </h2>
                    </div>
                    <div class="card-body text-center">

                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">Identité</th>
                                    <th scope="col">Adresse postale</th>
                                    <th scope="col">Téléphone</th>
                                    <th scope="col">Email</th>
                                    <th scope="col">Age</th>
                                    <th scope="col">Nationalité</th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        Jason Aymérick ALOYAU
                                    </td>
                                    <td>
                                         93800 Epinay-sur-Seine (France)
                                    </td>
                                    <td>
                                        
                                    </td>
                                    <td>
                                        aymerick.aloyau@outlook.fr
                                    </td>
                                    <td>
                                        28 ans
                                    </td>
                                    <td>
                                        Français
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Informations personnelles -->

                        <br>

                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Compétences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Energie -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Energie
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- Type of machines -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Type d'appareils
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Panneaux photovoltaïques / Éolienne / Turbine à gaz 
                                                                / Fusion nucléaire / Fusion froide / 
                                                                Machines à énergie libre
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Type of machines -->

                                        <br>

                                        <!-- Mathematics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mathématique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Nombres complexes / Géométrie élémentaire du plan / 
                                                                Géométrie élémentaire de l'espace / Fonctions usuelles 
                                                                / Équations différentielles linéaires / 
                                                                Étude des courbes planes / Conique / 
                                                                Entiers naturels, ensembles finis, nombres 
                                                                / Corps de nombres réels / Suites de nombres réels 
                                                                / Fonctions d'un réel variable réelle à valeurs 
                                                                multiples / Dérivation de fonctions à valeurs réelles 
                                                                / Intégration de fonctions à valeurs réelles sur 
                                                                un segment / Développements limités / Propriétés 
                                                                métriques des arcs / Suites et fonctions à valeurs 
                                                                complexes / Notions sur les fonctions de deux 
                                                                variables réelles / Intégrales multiples / Algébrique 
                                                                structures / Arithmétique / Polynômes / Fractions 
                                                                rationnelles / Espaces vectoriels / Dimension des 
                                                                espaces vectoriels / Calcul matriciel / Groupe 
                                                                symétrique, déterminant / Produit scalaire, 
                                                                groupe orthogonal / Géométrie affine / 
                                                                Techniques de démonstration / Techniques d'algèbre / 
                                                                Techniques d'analyse / Analyse multivariable / 
                                                                Théorie des graphes / Théorie de l'information / 
                                                                Analyse des données
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mathematics -->

                                        <br>

                                        <!-- Physics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Physique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Thermodynamique / Electrotechnique / Electromagnétisme 
                                                                / Electronique / Thermo-électricité / Mécanique 
                                                                quantique / Physique nucléaire / Electricité / 
                                                                Physique des plasmas / Résonance magnétique nucléaire / 
                                                                Matériaux diélectriques / Mécanique des fluides / 
                                                                Oscillateurs harmoniques et signaux sinusoïdaux / 
                                                                Circuits linéaires en régime continu / Dispositions 
                                                                transitoires du premier ordre / Régime de transition 
                                                                du second ordre / Régime sinusoïdal forcé - résonance / 
                                                                Filtrage / Cinématique / Principe fondamental de la 
                                                                dynamique / Énergie mécanique / Mouvement dans un champ 
                                                                électrique et magnétique / Propagation du signal - 
                                                                notion d'ondes / Lois de Snell-Descartes - réflexion 
                                                                et réfraction / Lentilles minces sphériques / Moment 
                                                                cinétique - solide tournant autour d'un axe fixe / 
                                                                Forces centrales conservatrices / États de la matière 
                                                                / Machines thermiques / Champ magnétique - forces de 
                                                                Laplace - induction
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Physics -->

                                        <br>

                                        <!-- Management -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Management
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Management collaboratif / Management commercial / 
                                                                Management conseil / Management d'équipe / 
                                                                Management de l'innovation / Management de projet / 
                                                                Management de la qualité / Management des connaissances 
                                                                / Management des opérations / Management des ressources 
                                                                humaines / Management des risques / Management du 
                                                                changement / Management évènementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Management -->

                                        <br>

                                        <!-- Marketing -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Marketing
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Marketing commercial / Marketing B2B / Marketing 
                                                                digital / Marketing direct / Marketing opérationnel / 
                                                                Marketing relationnel / Marketing stratégique / 
                                                                Marketing événementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Marketing -->

                                        <br>

                                        <!-- Finance -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Finance
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Analyse financière / Finance d'entreprise / Finance 
                                                                internationale / Finance personnelle / Financement de 
                                                                projet / Forex
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Finance -->

                                        <br>

                                        <!-- Economy -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Economie
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Macroéconomie / Microéconomie / Agents économiques 
                                                                / Économie agricole / Économie circulaire / 
                                                                Consommation et épargne / Économie de l'éducation / 
                                                                Économie financière / Économie géographique / 
                                                                Économie industrielle / Économie globale / 
                                                                Économie monétaire
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Economy -->

                                        <br>

                                        <!-- International law -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Loi internationale
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Site internet : "www.legifrance.gouv.fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "www.lexadin.nl"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://europa.eu/european-union/law/find-legislation_fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://www.fedlex.admin.ch/fr/home"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://legislation.gov.ky/"
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- International law -->        
                                    </div>
                                </div>
                                <!-- Energy -->

                                <br>

                                <!-- Computing -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Informatique
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- preferred programming languages -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Langages de programmation préférés
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python / Java / JavaScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Preferred programming languages -->    

                                        <br>

                                        <!-- Office -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Bureautique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Microsoft Excel, Word, PowerPoint, Access, 
                                                                Outlook
                                                            </li>
                                                            <li>
                                                                Open Office Calc, Writer
                                                            </li>
                                                            <li>
                                                                Intellij pour Java, PyCharm pour Python
                                                            </li>
                                                            <li>
                                                                Logiciel Pencil pour faire des schémas simplifiés / 
                                                                Logiciel GanttProject pour la gestion de projet
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Office -->

                                        <br>

                                        <!-- Front-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Front-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                HTML markup language /
                                                                CSS /
                                                                JQuery Framework /
                                                                Django Framework /
                                                                MJML Framework
                                                            </li>
                                                            <li>
                                                                Visual Studio Software Code /
                                                                Angular Framework 7 /
                                                                Boostrap /
                                                                Leaflet.js
                                                            </li>
                                                            <li>
                                                                Vis.js / 
                                                                Artyom.js /
                                                                Flowchart.js
                                                            </li>
                                                            <li>
                                                                Tomcat server /
                                                                Apache server /
                                                                Nginx server
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Front-end -->

                                        <br>

                                        <!-- Back-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Back-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python /
                                                                Flask /
                                                                BeautifulSoup /
                                                                pdfkit /
                                                                xlswriter /
                                                                xlrd /
                                                                pywinauto
                                                            </li>
                                                            <li>
                                                                Java /
                                                                Maven / 
                                                                Gradle /
                                                                Spring /
                                                                Jsoup /
                                                                Selenium WebDriver
                                                            </li>
                                                            <li>
                                                                JavaScript /
                                                                Visual Basic /
                                                                C# /
                                                                TypeScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Back-end -->

                                        <br>

                                        <!-- Tests unitaires et qualité -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Tests unitaires et qualité
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Unittests with Python / 
                                                                JUnit with Java / 
                                                                Mockito with Java / 
                                                                WireMock with Java
                                                            </li>
                                                            <li>
                                                                SonarQube
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Tests unitaires et qualité -->

                                        <br>

                                        <!-- Systèmes de bases de données -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes de base de données
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                MySQL / 
                                                                MySQL Workbench
                                                            </li>
                                                            <li>
                                                                NeO4J, serveur de base de données orienté graphe  / 
                                                                MongoDB, serveur de base de données orienté document
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes de bases de données -->

                                        <br>

                                        <!-- Outils de déploiement -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        DevOps
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Docker / Git
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Outils de déploiement -->

                                        <br>

                                        <!-- Intelligence artificielle -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Intelligence artificielle
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Exploration de données / 
                                                                Gestion des connaissances / 
                                                                Automatisation des processus robotiques /
                                                                Cloud computing / Reconnaissance vocale / 
                                                                Vision par ordinateur
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Intelligence artificielle -->

                                        <br>

                                        <!-- Mobile -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mobile
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Android /
                                                                Kivy
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mobile -->

                                        <br>

                                        <!-- Systèmes embarqués -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes embarqués
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Raspberry Pi / 
                                                                Arduino Uno and Arduino Yun / 
                                                                OpenWRT OS
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes embarqués -->

                                        <br>

                                        <!-- Conception assistée par ordinateur -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Conception assistée par ordinateur
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                FreeCAD / 
                                                                Skidl /
                                                                KiCAD
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Conception assistée par ordinateur -->

                                        <br>

                                        <!-- Gestion de version -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion des versions
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Git / 
                                                                GitHub : https://github.com/Jay4C
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de version -->

                                        <br>

                                        <!-- Réseaux et Systèmes -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Réseaux et systèmes
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Windows / 
                                                                Linux
                                                            </li>
                                                            <li>
                                                                Protocoles informatiques [HTTP, HTTPS, TCP, IPv4, IPv6, 
                                                                SSH, FTP, SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, 
                                                                Modbus RTU]
                                                            </li>
                                                            <li>
                                                                MinimalModbus avec Python / 
                                                                Jlibmodbus avec Java
                                                            </li>
                                                            <li>
                                                                Routeur 4G / 
                                                                Port Forwarding / 
                                                                DynDNS / 
                                                                Dongle 3G
                                                            </li>
                                                            <li>
                                                                Wireshark / 
                                                                PuTTY / 
                                                                WinSCP /
                                                                Postman /
                                                                TeamViewer / 
                                                                VirtualBox /
                                                                Tor /
                                                                Angry IP Scanner
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Réseaux et Systèmes -->

                                        <br>

                                        <!-- Gestion de projet et communications -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion de projet et communication
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Brainstorming / 
                                                                FreeMind / 
                                                                Pencil / 
                                                                Diagrammes UML/ 
                                                                ArgoUML
                                                            </li>
                                                            <li>
                                                                GanttProject / 
                                                                Planification des étapes d'un projet
                                                            </li>
                                                            <li>
                                                                Management de projet / Management des opérations 
                                                                / Management agile / Management de la qualité
                                                            </li>
                                                            <li>
                                                                Analyse des besoins client / 
                                                                Etablissement d'un cahier des charges
                                                            </li>
                                                            <li>
                                                                Rédaction d'un support technique / 
                                                                Rédaction d'une base de connaissances
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de projet et communications -->        
                                    </div>
                                </div>
                                <!-- Computing -->
                            </div>
                        </div>
                        <!-- Compétences -->

                        <br>

                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Société PASS Technologie -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Développeur informatique polyvalent
                                                <br>
                                                Période : février 2019 - avril 2019 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Société : Pass Technology - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application pour le BIM (Building 
                                                        Information Modeling)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Pass Tech -->

                                <br>

                                <!-- Kriir -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Ingénieur informatique
                                                <br>
                                                Période : septembre - novembre 2018 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Entreprise : Kriir - Ermont
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application web Java pour utiliser la 
                                                        signature électronique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Kriir -->

                                <br>

                                <!-- Refuel -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Lire et écrire sur un automate en Modbus
                                                <br>
                                                Période : juin - juillet 2018 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Refuel S.A.S. - Trappes
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de plusieurs fonctionnalités pour lire et écrire 
                                                        sur un automate programmable via le protocole Modbus TCP avec 
                                                        langage Java en liaison WiFi
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Refuel  -->

                                <br>

                                <!-- SETEC SMART EFFICIENCY -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Génération automatique de rapport énergétique
                                                <br>
                                                Période : novembre - décembre 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : SETEC SMART EFFICIENCY - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développer un programme en Java pour générer des rapports 
                                                        énergétique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SETEC SMART EFFICIENCY -->

                                <br>

                                <!--  SOCIÉTÉ GAMELOFT -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Vérification de la pertinence des données
                                                <br>
                                                Période : juin - juillet 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Gameloft - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une macro Visual Basic pour un fichier 
                                                        Excel permettant de contrôler la cohérence des données saisies 
                                                        dans plusieurs fichiers Excel et de consolider ces données 
                                                        dans un seul fichier Excel pour les ressources humaines
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SOCIÉTÉ GAMELOFT -->
                            </div>
                        </div>
                        <!-- Expériences -->

                        <br>

                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Formation Java -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Formation Java / JEE
                                                <br>
                                                Période : avril - juillet 2017 [4 mois]
                                                <br>
                                                Centre de formation : INTI Formation - Paris Tour Montparnasse
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Apprentissage du langage informatique Java / 
                                                        Apprentissage de la programmation orientée objet
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Formation Java -->

                                <br>

                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : septembre 2013 - août 2016 [3 ans]
                                                <br>
                                                Centre de formation : Ecole Supérieure d'Ingénieurs Léonard de Vinci 
                                                - Paris La Défense (France)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / Informatique / Finance / Mécanique 
                                                        numérique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->

                        <br>

                        <!-- Languages -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Langues
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Anglais écrit et oral [TOEIC : 840 points ; Date d'expiration : 11 juillet 2020]
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Languages -->

                        <br>

                        <!-- Centre d'intérêts -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Centre d'intérêts
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Technologie / Brevets
                                    </li>
                                    <li>
                                        Droit français / 
                                        Droit européen / 
                                        Droit international
                                    </li>
                                    <li>
                                        Physique / 
                                        Mathématiques / 
                                        Finance / 
                                        Economie / 
                                        Gestion / 
                                        Marketing / 
                                        Spiritualité / 
                                        Ufologie
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Centre d'intérêts -->
                    </div>
                </div>
            </div>

            <br>

            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>
        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Ingénieur en énergies renouvelables'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body,
                           "CV\\CV__[Ingenieur_Energies_Renouvelables]_v1.pdf", configuration=config,
                           options=options
                           )

    # ok
    def test_cv_ingenieur_energie_renouvelable_v2(self):
        print("test_cv_ingenieur_energie_renouvelable_v2")

        body = """
        <!doctype html>
        <html lang="en">

        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

            <title>CV de  - Ingénieur en énergies renouvelables</title>
        </head>

        <body>
            <br>
            <div class="container">
                <div class="card text-center">
                    <div class="card-header text-center">
                        <h2>
                            Ingénieur en énergies renouvelables
                        </h2>
                    </div>
                    <div class="card-body text-center">

                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">Identité</th>
                                    <th scope="col">Adresse postale</th>
                                    <th scope="col">Téléphone</th>
                                    <th scope="col">Email</th>
                                    <th scope="col">Age</th>
                                    <th scope="col">Nationalité</th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        Jason Aymérick ALOYAU
                                    </td>
                                    <td>
                                         93800 Epinay-sur-Seine (France)
                                    </td>
                                    <td>
                                        
                                    </td>
                                    <td>
                                        jason.aymerick.aloyau@gmail.com
                                    </td>
                                    <td>
                                        28 ans
                                    </td>
                                    <td>
                                        Français
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Informations personnelles -->

                        <br>

                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Compétences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Energie -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Energie
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- Type of machines -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Type d'appareils
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Panneaux photovoltaïques / Éolienne / Turbine à gaz 
                                                                / Fusion nucléaire / Fusion froide / 
                                                                Machines à énergie libre
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Type of machines -->

                                        <br>

                                        <!-- Mathematics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mathématique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Nombres complexes / Géométrie élémentaire du plan / 
                                                                Géométrie élémentaire de l'espace / Fonctions usuelles 
                                                                / Équations différentielles linéaires / 
                                                                Étude des courbes planes / Conique / 
                                                                Entiers naturels, ensembles finis, nombres 
                                                                / Corps de nombres réels / Suites de nombres réels 
                                                                / Fonctions d'un réel variable réelle à valeurs 
                                                                multiples / Dérivation de fonctions à valeurs réelles 
                                                                / Intégration de fonctions à valeurs réelles sur 
                                                                un segment / Développements limités / Propriétés 
                                                                métriques des arcs / Suites et fonctions à valeurs 
                                                                complexes / Notions sur les fonctions de deux 
                                                                variables réelles / Intégrales multiples / Algébrique 
                                                                structures / Arithmétique / Polynômes / Fractions 
                                                                rationnelles / Espaces vectoriels / Dimension des 
                                                                espaces vectoriels / Calcul matriciel / Groupe 
                                                                symétrique, déterminant / Produit scalaire, 
                                                                groupe orthogonal / Géométrie affine / 
                                                                Techniques de démonstration / Techniques d'algèbre / 
                                                                Techniques d'analyse / Analyse multivariable / 
                                                                Théorie des graphes / Théorie de l'information / 
                                                                Analyse des données
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mathematics -->

                                        <br>

                                        <!-- Physics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Physique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Thermodynamique / Electrotechnique / Electromagnétisme 
                                                                / Electronique / Thermo-électricité / Mécanique 
                                                                quantique / Physique nucléaire / Electricité / 
                                                                Physique des plasmas / Résonance magnétique nucléaire / 
                                                                Matériaux diélectriques / Mécanique des fluides / 
                                                                Oscillateurs harmoniques et signaux sinusoïdaux / 
                                                                Circuits linéaires en régime continu / Dispositions 
                                                                transitoires du premier ordre / Régime de transition 
                                                                du second ordre / Régime sinusoïdal forcé - résonance / 
                                                                Filtrage / Cinématique / Principe fondamental de la 
                                                                dynamique / Énergie mécanique / Mouvement dans un champ 
                                                                électrique et magnétique / Propagation du signal - 
                                                                notion d'ondes / Lois de Snell-Descartes - réflexion 
                                                                et réfraction / Lentilles minces sphériques / Moment 
                                                                cinétique - solide tournant autour d'un axe fixe / 
                                                                Forces centrales conservatrices / États de la matière 
                                                                / Machines thermiques / Champ magnétique - forces de 
                                                                Laplace - induction
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Physics -->

                                        <br>

                                        <!-- Management -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Management
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Management collaboratif / Management commercial / 
                                                                Management conseil / Management d'équipe / 
                                                                Management de l'innovation / Management de projet / 
                                                                Management de la qualité / Management des connaissances 
                                                                / Management des opérations / Management des ressources 
                                                                humaines / Management des risques / Management du 
                                                                changement / Management évènementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Management -->

                                        <br>

                                        <!-- Marketing -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Marketing
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Marketing commercial / Marketing B2B / Marketing 
                                                                digital / Marketing direct / Marketing opérationnel / 
                                                                Marketing relationnel / Marketing stratégique / 
                                                                Marketing événementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Marketing -->

                                        <br>

                                        <!-- Finance -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Finance
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Analyse financière / Finance d'entreprise / Finance 
                                                                internationale / Finance personnelle / Financement de 
                                                                projet / Forex
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Finance -->

                                        <br>

                                        <!-- Economy -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Economie
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Macroéconomie / Microéconomie / Agents économiques 
                                                                / Économie agricole / Économie circulaire / 
                                                                Consommation et épargne / Économie de l'éducation / 
                                                                Économie financière / Économie géographique / 
                                                                Économie industrielle / Économie globale / 
                                                                Économie monétaire
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Economy -->

                                        <br>

                                        <!-- International law -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Loi internationale
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Site internet : "www.legifrance.gouv.fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "www.lexadin.nl"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://europa.eu/european-union/law/find-legislation_fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://www.fedlex.admin.ch/fr/home"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://legislation.gov.ky/"
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- International law -->        
                                    </div>
                                </div>
                                <!-- Energy -->

                                <br>

                                <!-- Computing -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Informatique
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- preferred programming languages -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Langages de programmation préférés
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python / Java / JavaScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Preferred programming languages -->    

                                        <br>

                                        <!-- Office -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Bureautique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Microsoft Excel, Word, PowerPoint, Access, 
                                                                Outlook
                                                            </li>
                                                            <li>
                                                                Open Office Calc, Writer
                                                            </li>
                                                            <li>
                                                                Intellij pour Java, PyCharm pour Python
                                                            </li>
                                                            <li>
                                                                Logiciel Pencil pour faire des schémas simplifiés / 
                                                                Logiciel GanttProject pour la gestion de projet
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Office -->

                                        <br>

                                        <!-- Front-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Front-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                HTML markup language /
                                                                CSS /
                                                                JQuery Framework /
                                                                Django Framework /
                                                                MJML Framework
                                                            </li>
                                                            <li>
                                                                Visual Studio Software Code /
                                                                Angular Framework 7 /
                                                                Boostrap /
                                                                Leaflet.js
                                                            </li>
                                                            <li>
                                                                Vis.js / 
                                                                Artyom.js /
                                                                Flowchart.js
                                                            </li>
                                                            <li>
                                                                Tomcat server /
                                                                Apache server /
                                                                Nginx server
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Front-end -->

                                        <br>

                                        <!-- Back-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Back-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python /
                                                                Flask /
                                                                BeautifulSoup /
                                                                pdfkit /
                                                                xlswriter /
                                                                xlrd /
                                                                pywinauto
                                                            </li>
                                                            <li>
                                                                Java /
                                                                Maven / 
                                                                Gradle /
                                                                Spring /
                                                                Jsoup /
                                                                Selenium WebDriver
                                                            </li>
                                                            <li>
                                                                JavaScript /
                                                                Visual Basic /
                                                                C# /
                                                                TypeScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Back-end -->

                                        <br>

                                        <!-- Tests unitaires et qualité -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Tests unitaires et qualité
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Unittests with Python / 
                                                                JUnit with Java / 
                                                                Mockito with Java / 
                                                                WireMock with Java
                                                            </li>
                                                            <li>
                                                                SonarQube
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Tests unitaires et qualité -->

                                        <br>

                                        <!-- Systèmes de bases de données -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes de base de données
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                MySQL / 
                                                                MySQL Workbench
                                                            </li>
                                                            <li>
                                                                NeO4J, serveur de base de données orienté graphe  / 
                                                                MongoDB, serveur de base de données orienté document
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes de bases de données -->

                                        <br>

                                        <!-- Outils de déploiement -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        DevOps
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Docker / Git
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Outils de déploiement -->

                                        <br>

                                        <!-- Intelligence artificielle -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Intelligence artificielle
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Exploration de données / 
                                                                Gestion des connaissances / 
                                                                Automatisation des processus robotiques /
                                                                Cloud computing / Reconnaissance vocale / 
                                                                Vision par ordinateur
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Intelligence artificielle -->

                                        <br>

                                        <!-- Mobile -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mobile
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Android /
                                                                Kivy
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mobile -->

                                        <br>

                                        <!-- Systèmes embarqués -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes embarqués
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Raspberry Pi / 
                                                                Arduino Uno and Arduino Yun / 
                                                                OpenWRT OS
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes embarqués -->

                                        <br>

                                        <!-- Conception assistée par ordinateur -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Conception assistée par ordinateur
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                FreeCAD / 
                                                                Skidl /
                                                                KiCAD
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Conception assistée par ordinateur -->

                                        <br>

                                        <!-- Gestion de version -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion des versions
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Git / 
                                                                GitHub : https://github.com/Jay4C
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de version -->

                                        <br>

                                        <!-- Réseaux et Systèmes -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Réseaux et systèmes
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Windows / 
                                                                Linux
                                                            </li>
                                                            <li>
                                                                Protocoles informatiques [HTTP, HTTPS, TCP, IPv4, IPv6, 
                                                                SSH, FTP, SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, 
                                                                Modbus RTU]
                                                            </li>
                                                            <li>
                                                                MinimalModbus avec Python / 
                                                                Jlibmodbus avec Java
                                                            </li>
                                                            <li>
                                                                Routeur 4G / 
                                                                Port Forwarding / 
                                                                DynDNS / 
                                                                Dongle 3G
                                                            </li>
                                                            <li>
                                                                Wireshark / 
                                                                PuTTY / 
                                                                WinSCP /
                                                                Postman /
                                                                TeamViewer / 
                                                                VirtualBox /
                                                                Tor /
                                                                Angry IP Scanner
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Réseaux et Systèmes -->

                                        <br>

                                        <!-- Gestion de projet et communications -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion de projet et communication
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Brainstorming / 
                                                                FreeMind / 
                                                                Pencil / 
                                                                Diagrammes UML/ 
                                                                ArgoUML
                                                            </li>
                                                            <li>
                                                                GanttProject / 
                                                                Planification des étapes d'un projet
                                                            </li>
                                                            <li>
                                                                Management de projet / Management des opérations 
                                                                / Management agile / Management de la qualité
                                                            </li>
                                                            <li>
                                                                Analyse des besoins client / 
                                                                Etablissement d'un cahier des charges
                                                            </li>
                                                            <li>
                                                                Rédaction d'un support technique / 
                                                                Rédaction d'une base de connaissances
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de projet et communications -->        
                                    </div>
                                </div>
                                <!-- Computing -->
                            </div>
                        </div>
                        <!-- Compétences -->

                        <br>

                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Société PASS Technologie -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Développeur informatique polyvalent
                                                <br>
                                                Période : février 2019 - avril 2019 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Société : Pass Technology - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application pour le BIM (Building 
                                                        Information Modeling)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Pass Tech -->

                                <br>

                                <!-- Kriir -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Ingénieur informatique
                                                <br>
                                                Période : septembre - novembre 2018 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Entreprise : Kriir - Ermont
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application web Java pour utiliser la 
                                                        signature électronique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Kriir -->

                                <br>

                                <!-- Refuel -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Lire et écrire sur un automate en Modbus
                                                <br>
                                                Période : juin - juillet 2018 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Refuel S.A.S. - Trappes
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de plusieurs fonctionnalités pour lire et écrire 
                                                        sur un automate programmable via le protocole Modbus TCP avec 
                                                        langage Java en liaison WiFi
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Refuel  -->

                                <br>

                                <!-- SETEC SMART EFFICIENCY -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Génération automatique de rapport énergétique
                                                <br>
                                                Période : novembre - décembre 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : SETEC SMART EFFICIENCY - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développer un programme en Java pour générer des rapports 
                                                        énergétique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SETEC SMART EFFICIENCY -->

                                <br>

                                <!--  SOCIÉTÉ GAMELOFT -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Vérification de la pertinence des données
                                                <br>
                                                Période : juin - juillet 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Gameloft - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une macro Visual Basic pour un fichier 
                                                        Excel permettant de contrôler la cohérence des données saisies 
                                                        dans plusieurs fichiers Excel et de consolider ces données 
                                                        dans un seul fichier Excel pour les ressources humaines
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SOCIÉTÉ GAMELOFT -->
                            </div>
                        </div>
                        <!-- Expériences -->

                        <br>

                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Formation Java -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Formation Java / JEE
                                                <br>
                                                Période : avril - juillet 2017 [4 mois]
                                                <br>
                                                Centre de formation : INTI Formation - Paris Tour Montparnasse
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Apprentissage du langage informatique Java / 
                                                        Apprentissage de la programmation orientée objet
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Formation Java -->

                                <br>

                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : septembre 2013 - août 2016 [3 ans]
                                                <br>
                                                Centre de formation : Ecole Supérieure d'Ingénieurs Léonard de Vinci 
                                                - Paris La Défense (France)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / Informatique / Finance / Mécanique 
                                                        numérique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->

                        <br>

                        <!-- Languages -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Langues
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Anglais écrit et oral [TOEIC : 840 points ; Date d'expiration : 11 juillet 2020]
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Languages -->

                        <br>

                        <!-- Centre d'intérêts -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Centre d'intérêts
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Technologie / Brevets
                                    </li>
                                    <li>
                                        Droit français / 
                                        Droit européen / 
                                        Droit international
                                    </li>
                                    <li>
                                        Physique / 
                                        Mathématiques / 
                                        Finance / 
                                        Economie / 
                                        Gestion / 
                                        Marketing / 
                                        Spiritualité / 
                                        Ufologie
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Centre d'intérêts -->
                    </div>
                </div>
            </div>

            <br>

            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>
        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Ingénieur en énergies renouvelables'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body,
                           "CV\\CV__[Ingenieur_Energies_Renouvelables]_v2.pdf",
                           configuration=config,
                           options=options
                           )

    # ok
    def test_cv_ingenieur_energie_renouvelable_v3(self):
        print("test_cv_ingenieur_energie_renouvelable_v3")

        body = """
        <!doctype html>
        <html lang="en">

        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

            <title>CV de  - Ingénieur en énergies renouvelables</title>
        </head>

        <body>
            <br>
            <div class="container">
                <div class="card text-center">
                    <div class="card-header text-center">
                        <h2>
                            Ingénieur en énergies renouvelables
                        </h2>
                    </div>
                    <div class="card-body text-center">

                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">Identité</th>
                                    <th scope="col">Adresse postale</th>
                                    <th scope="col">Téléphone</th>
                                    <th scope="col">Email</th>
                                    <th scope="col">Age</th>
                                    <th scope="col">Nationalité</th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        Jason Aymérick ALOYAU
                                    </td>
                                    <td>
                                         93800 Epinay-sur-Seine (France)
                                    </td>
                                    <td>
                                        
                                    </td>
                                    <td>
                                        
                                    </td>
                                    <td>
                                        28 ans
                                    </td>
                                    <td>
                                        Français
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Informations personnelles -->

                        <br>

                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Compétences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Energie -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Energie
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- Type of machines -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Type d'appareils
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Panneaux photovoltaïques / Éolienne / Turbine à gaz 
                                                                / Fusion nucléaire / Fusion froide / 
                                                                Machines à énergie libre
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Type of machines -->

                                        <br>

                                        <!-- Mathematics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mathématique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Nombres complexes / Géométrie élémentaire du plan / 
                                                                Géométrie élémentaire de l'espace / Fonctions usuelles 
                                                                / Équations différentielles linéaires / 
                                                                Étude des courbes planes / Conique / 
                                                                Entiers naturels, ensembles finis, nombres 
                                                                / Corps de nombres réels / Suites de nombres réels 
                                                                / Fonctions d'un réel variable réelle à valeurs 
                                                                multiples / Dérivation de fonctions à valeurs réelles 
                                                                / Intégration de fonctions à valeurs réelles sur 
                                                                un segment / Développements limités / Propriétés 
                                                                métriques des arcs / Suites et fonctions à valeurs 
                                                                complexes / Notions sur les fonctions de deux 
                                                                variables réelles / Intégrales multiples / Algébrique 
                                                                structures / Arithmétique / Polynômes / Fractions 
                                                                rationnelles / Espaces vectoriels / Dimension des 
                                                                espaces vectoriels / Calcul matriciel / Groupe 
                                                                symétrique, déterminant / Produit scalaire, 
                                                                groupe orthogonal / Géométrie affine / 
                                                                Techniques de démonstration / Techniques d'algèbre / 
                                                                Techniques d'analyse / Analyse multivariable / 
                                                                Théorie des graphes / Théorie de l'information / 
                                                                Analyse des données
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mathematics -->

                                        <br>

                                        <!-- Physics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Physique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Thermodynamique / Electrotechnique / Electromagnétisme 
                                                                / Electronique / Thermo-électricité / Mécanique 
                                                                quantique / Physique nucléaire / Electricité / 
                                                                Physique des plasmas / Résonance magnétique nucléaire / 
                                                                Matériaux diélectriques / Mécanique des fluides / 
                                                                Oscillateurs harmoniques et signaux sinusoïdaux / 
                                                                Circuits linéaires en régime continu / Dispositions 
                                                                transitoires du premier ordre / Régime de transition 
                                                                du second ordre / Régime sinusoïdal forcé - résonance / 
                                                                Filtrage / Cinématique / Principe fondamental de la 
                                                                dynamique / Énergie mécanique / Mouvement dans un champ 
                                                                électrique et magnétique / Propagation du signal - 
                                                                notion d'ondes / Lois de Snell-Descartes - réflexion 
                                                                et réfraction / Lentilles minces sphériques / Moment 
                                                                cinétique - solide tournant autour d'un axe fixe / 
                                                                Forces centrales conservatrices / États de la matière 
                                                                / Machines thermiques / Champ magnétique - forces de 
                                                                Laplace - induction
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Physics -->

                                        <br>

                                        <!-- Management -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Management
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Management collaboratif / Management commercial / 
                                                                Management conseil / Management d'équipe / 
                                                                Management de l'innovation / Management de projet / 
                                                                Management de la qualité / Management des connaissances 
                                                                / Management des opérations / Management des ressources 
                                                                humaines / Management des risques / Management du 
                                                                changement / Management évènementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Management -->

                                        <br>

                                        <!-- Marketing -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Marketing
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Marketing commercial / Marketing B2B / Marketing 
                                                                digital / Marketing direct / Marketing opérationnel / 
                                                                Marketing relationnel / Marketing stratégique / 
                                                                Marketing événementiel
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Marketing -->

                                        <br>

                                        <!-- Finance -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Finance
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Analyse financière / Finance d'entreprise / Finance 
                                                                internationale / Finance personnelle / Financement de 
                                                                projet / Forex
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Finance -->

                                        <br>

                                        <!-- Economy -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Economie
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Macroéconomie / Microéconomie / Agents économiques 
                                                                / Économie agricole / Économie circulaire / 
                                                                Consommation et épargne / Économie de l'éducation / 
                                                                Économie financière / Économie géographique / 
                                                                Économie industrielle / Économie globale / 
                                                                Économie monétaire
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Economy -->

                                        <br>

                                        <!-- International law -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Loi internationale
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Site internet : "www.legifrance.gouv.fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "www.lexadin.nl"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://europa.eu/european-union/law/find-legislation_fr"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://www.fedlex.admin.ch/fr/home"
                                                            </li>
                                                            <li>
                                                                Site internet : "https://legislation.gov.ky/"
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- International law -->        
                                    </div>
                                </div>
                                <!-- Energy -->

                                <br>

                                <!-- Computing -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Informatique
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- preferred programming languages -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Langages de programmation préférés
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python / Java / JavaScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Preferred programming languages -->    

                                        <br>

                                        <!-- Office -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Bureautique
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Microsoft Excel, Word, PowerPoint, Access, 
                                                                Outlook
                                                            </li>
                                                            <li>
                                                                Open Office Calc, Writer
                                                            </li>
                                                            <li>
                                                                Intellij pour Java, PyCharm pour Python
                                                            </li>
                                                            <li>
                                                                Logiciel Pencil pour faire des schémas simplifiés / 
                                                                Logiciel GanttProject pour la gestion de projet
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Office -->

                                        <br>

                                        <!-- Front-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Front-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                HTML markup language /
                                                                CSS /
                                                                JQuery Framework /
                                                                Django Framework /
                                                                MJML Framework
                                                            </li>
                                                            <li>
                                                                Visual Studio Software Code /
                                                                Angular Framework 7 /
                                                                Boostrap /
                                                                Leaflet.js
                                                            </li>
                                                            <li>
                                                                Vis.js / 
                                                                Artyom.js /
                                                                Flowchart.js
                                                            </li>
                                                            <li>
                                                                Tomcat server /
                                                                Apache server /
                                                                Nginx server
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Front-end -->

                                        <br>

                                        <!-- Back-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Back-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python /
                                                                Flask /
                                                                BeautifulSoup /
                                                                pdfkit /
                                                                xlswriter /
                                                                xlrd /
                                                                pywinauto
                                                            </li>
                                                            <li>
                                                                Java /
                                                                Maven / 
                                                                Gradle /
                                                                Spring /
                                                                Jsoup /
                                                                Selenium WebDriver
                                                            </li>
                                                            <li>
                                                                JavaScript /
                                                                Visual Basic /
                                                                C# /
                                                                TypeScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Back-end -->

                                        <br>

                                        <!-- Tests unitaires et qualité -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Tests unitaires et qualité
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Unittests with Python / 
                                                                JUnit with Java / 
                                                                Mockito with Java / 
                                                                WireMock with Java
                                                            </li>
                                                            <li>
                                                                SonarQube
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Tests unitaires et qualité -->

                                        <br>

                                        <!-- Systèmes de bases de données -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes de base de données
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                MySQL / 
                                                                MySQL Workbench
                                                            </li>
                                                            <li>
                                                                NeO4J, serveur de base de données orienté graphe  / 
                                                                MongoDB, serveur de base de données orienté document
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes de bases de données -->

                                        <br>

                                        <!-- Outils de déploiement -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        DevOps
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Docker / Git
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Outils de déploiement -->

                                        <br>

                                        <!-- Intelligence artificielle -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Intelligence artificielle
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Exploration de données / 
                                                                Gestion des connaissances / 
                                                                Automatisation des processus robotiques /
                                                                Cloud computing / Reconnaissance vocale / 
                                                                Vision par ordinateur
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Intelligence artificielle -->

                                        <br>

                                        <!-- Mobile -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mobile
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Android /
                                                                Kivy
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mobile -->

                                        <br>

                                        <!-- Systèmes embarqués -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Systèmes embarqués
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Raspberry Pi / 
                                                                Arduino Uno and Arduino Yun / 
                                                                OpenWRT OS
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes embarqués -->

                                        <br>

                                        <!-- Conception assistée par ordinateur -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Conception assistée par ordinateur
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                FreeCAD / 
                                                                Skidl /
                                                                KiCAD
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Conception assistée par ordinateur -->

                                        <br>

                                        <!-- Gestion de version -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion des versions
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Git / 
                                                                GitHub : https://github.com/Jay4C
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de version -->

                                        <br>

                                        <!-- Réseaux et Systèmes -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Réseaux et systèmes
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Windows / 
                                                                Linux
                                                            </li>
                                                            <li>
                                                                Protocoles informatiques [HTTP, HTTPS, TCP, IPv4, IPv6, 
                                                                SSH, FTP, SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, 
                                                                Modbus RTU]
                                                            </li>
                                                            <li>
                                                                MinimalModbus avec Python / 
                                                                Jlibmodbus avec Java
                                                            </li>
                                                            <li>
                                                                Routeur 4G / 
                                                                Port Forwarding / 
                                                                DynDNS / 
                                                                Dongle 3G
                                                            </li>
                                                            <li>
                                                                Wireshark / 
                                                                PuTTY / 
                                                                WinSCP /
                                                                Postman /
                                                                TeamViewer / 
                                                                VirtualBox /
                                                                Tor /
                                                                Angry IP Scanner
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Réseaux et Systèmes -->

                                        <br>

                                        <!-- Gestion de projet et communications -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Gestion de projet et communication
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Brainstorming / 
                                                                FreeMind / 
                                                                Pencil / 
                                                                Diagrammes UML/ 
                                                                ArgoUML
                                                            </li>
                                                            <li>
                                                                GanttProject / 
                                                                Planification des étapes d'un projet
                                                            </li>
                                                            <li>
                                                                Management de projet / Management des opérations 
                                                                / Management agile / Management de la qualité
                                                            </li>
                                                            <li>
                                                                Analyse des besoins client / 
                                                                Etablissement d'un cahier des charges
                                                            </li>
                                                            <li>
                                                                Rédaction d'un support technique / 
                                                                Rédaction d'une base de connaissances
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de projet et communications -->        
                                    </div>
                                </div>
                                <!-- Computing -->
                            </div>
                        </div>
                        <!-- Compétences -->

                        <br>

                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Expériences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Société PASS Technologie -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Développeur informatique polyvalent
                                                <br>
                                                Période : février 2019 - avril 2019 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Société : Pass Technology - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application pour le BIM (Building 
                                                        Information Modeling)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Pass Tech -->

                                <br>

                                <!-- Kriir -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Ingénieur informatique
                                                <br>
                                                Période : septembre - novembre 2018 [3 mois]
                                                <br>
                                                Statut professionnel : salarié
                                                <br>
                                                Entreprise : Kriir - Ermont
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une application web Java pour utiliser la 
                                                        signature électronique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Kriir -->

                                <br>

                                <!-- Refuel -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Lire et écrire sur un automate en Modbus
                                                <br>
                                                Période : juin - juillet 2018 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Refuel S.A.S. - Trappes
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement de plusieurs fonctionnalités pour lire et écrire 
                                                        sur un automate programmable via le protocole Modbus TCP avec 
                                                        langage Java en liaison WiFi
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Refuel  -->

                                <br>

                                <!-- SETEC SMART EFFICIENCY -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Génération automatique de rapport énergétique
                                                <br>
                                                Période : novembre - décembre 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : SETEC SMART EFFICIENCY - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développer un programme en Java pour générer des rapports 
                                                        énergétique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SETEC SMART EFFICIENCY -->

                                <br>

                                <!--  SOCIÉTÉ GAMELOFT -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Vérification de la pertinence des données
                                                <br>
                                                Période : juin - juillet 2017 [2 mois]
                                                <br>
                                                Statut professionnel : Freelance
                                                <br>
                                                Société : Gameloft - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Développement d'une macro Visual Basic pour un fichier 
                                                        Excel permettant de contrôler la cohérence des données saisies 
                                                        dans plusieurs fichiers Excel et de consolider ces données 
                                                        dans un seul fichier Excel pour les ressources humaines
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SOCIÉTÉ GAMELOFT -->
                            </div>
                        </div>
                        <!-- Expériences -->

                        <br>

                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Formations
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Formation Java -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Formation Java / JEE
                                                <br>
                                                Période : avril - juillet 2017 [4 mois]
                                                <br>
                                                Centre de formation : INTI Formation - Paris Tour Montparnasse
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Apprentissage du langage informatique Java / 
                                                        Apprentissage de la programmation orientée objet
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Formation Java -->

                                <br>

                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Titre : Diplôme d'ingénieur généraliste
                                                <br>
                                                Période : septembre 2013 - août 2016 [3 ans]
                                                <br>
                                                Centre de formation : Ecole Supérieure d'Ingénieurs Léonard de Vinci 
                                                - Paris La Défense (France)
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Nouvelles énergies / Informatique / Finance / Mécanique 
                                                        numérique
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->

                        <br>

                        <!-- Languages -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Langues
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Anglais écrit et oral [TOEIC : 840 points ; Date d'expiration : 11 juillet 2020]
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Languages -->

                        <br>

                        <!-- Centre d'intérêts -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Centre d'intérêts
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Technologie / Brevets
                                    </li>
                                    <li>
                                        Droit français / 
                                        Droit européen / 
                                        Droit international
                                    </li>
                                    <li>
                                        Physique / 
                                        Mathématiques / 
                                        Finance / 
                                        Economie / 
                                        Gestion / 
                                        Marketing / 
                                        Spiritualité / 
                                        Ufologie
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Centre d'intérêts -->
                    </div>
                </div>
            </div>

            <br>

            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>
        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'CV de ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Ingénieur en énergies renouvelables'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body,
                           "CV\\CV__[Ingenieur_Energies_Renouvelables]_v3.pdf",
                           configuration=config,
                           options=options
                           )


class Resume(unittest.TestCase):
    # ok
    def test_resume_full_stack_developper(self):
        print("test_resume")

        body = """
    <!doctype html>
    <html lang="en">

    <head>
        <!-- Required meta tags -->
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">

        <!-- Bootstrap CSS -->
        <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
            integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">

        <title>Resume of Mr.  - Full-Stack Developer</title>
    </head>

    <body>
        <br>
        <div class="container">
            <div class="card text-center">
                <div class="card-header text-center">
                    <h2>
                        Full-Stack Developer
                    </h2>
                </div>
                <div class="card-body text-center">

                    <!-- Informations personnelles -->
                    <table class="table table-striped table-bordered">
                        <thead>
                            <tr>
                                <th scope="col">Identity</th>
                                <th scope="col">Postal address</th>
                                <th scope="col">Phone</th>
                                <th scope="col">Email</th>
                                <th scope="col">Nationality</th>
                            </tr>
                        </thead>
                        <tbody>
                            <tr>
                                <td>
                                    Mr. 
                                </td>
                                <td>
                                    31 Avenue de Ségur 75007 Paris
                                </td>
                                <td></td>
                                <td>@holomorphe.com</td>
                                <td>French</td>
                            </tr>
                        </tbody>
                    </table>

                    <!--  Compétences -->
                    <div class="card">
                        <div class="card-header">
                            <h4>
                                Skills
                            </h4>
                        </div>
                        <div class="card-body">
                            <!-- Anglais -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            English
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    English written and oral [TOEIC: 840 points; Expiry date: July 11, 2020]
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Anglais -->

                            <br>

                            <!-- Langages de programmation de prédilection -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Predilection programming languages
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>Python / Java / JavaScript</li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Langages de programmation de prédilection -->

                            <br>

                            <!-- Bureautique -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Office
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>Microsoft Excel, Word, PowerPoint, Access, Outlook Software</li>
                                                <li>Open Office Calc, Writer Software</li>
                                                <li>
                                                    Intellij Software for Java / 
                                                    Python Python Software
                                                </li>
                                                <li>
                                                    Pencil software to make simplified schemes / 
                                                    GanttProject software for project management
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Bureautique -->

                            <br>

                            <!-- Front-end -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Front-end
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    HTML markup language /
                                                    CSS /
                                                    JQuery Framework /
                                                    Django Framework /
                                                    MJML Framework for Email Responsive
                                                </li>
                                                <li>
                                                    Visual Studio Software Code /
                                                    Angular Framework 7 /
                                                    Boostrap /
                                                    Leaflet.js
                                                </li>
                                                <li>
                                                    Vis.js / 
                                                    Artyom.js /
                                                    Flowchart.js
                                                </li>
                                                <li>
                                                    Tomcat server /
                                                    Apache server /
                                                    Nginx server
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Front-end -->

                            <br>

                            <!-- Back-end -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Back-end
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Python /
                                                    Flask /
                                                    BeautifulSoup /
                                                    pdfkit /
                                                    xlswriter /
                                                    xlrd /
                                                    pywinauto /
                                                    Hyperledger Sawtooth
                                                </li>
                                                <li>
                                                    Java /
                                                    Maven / 
                                                    Gradle /
                                                    Spring /
                                                    Jsoup /
                                                    Selenium WebDriver
                                                </li>
                                                <li>
                                                    JavaScript /
                                                    Visual Basic /
                                                    C# /
                                                    TypeScript
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Back-end -->

                            <br>

                            <!-- Tests unitaires et qualité -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Unit and quality tests
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Unittests with Python / 
                                                    JUnit with Java / 
                                                    Mockito with Java / 
                                                    WireMock with Java
                                                </li>
                                                <li>
                                                    SonarQube
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Tests unitaires et qualité -->

                            <br>

                            <!-- Systèmes de bases de données -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Database systems
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    MySQL / 
                                                    MySQL Workbench
                                                </li>
                                                <li>
                                                    NeO4J, graph database server  / 
                                                    MongoDB, documents database server
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Systèmes de bases de données -->

                            <br>

                            <!-- Outils de déploiement -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            DevOps
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Docker / Git
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Outils de déploiement -->

                            <br>

                            <!-- Intelligence artificielle -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Artificial intelligence
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Data mining / 
                                                    Knowledge management /
                                                    Robotic Process Automation
                                                </li>
                                                <li>
                                                    Cloud computing [Open data; Autodesk Forge; Netheos] / 
                                                    Speech recognition / 
                                                    Computer vision
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Intelligence artificielle -->

                            <br>

                            <!-- Mobile -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Mobile
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Android /
                                                    Kivy
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Mobile -->

                            <br>

                            <!-- Systèmes embarqués -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Embedded systems
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Raspberry Pi
                                                </li>
                                                <li>
                                                    Arduino Uno and Arduino Yun /
                                                    OpenWRT OS
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Systèmes embarqués -->

                            <br>

                            <!-- Conception assistée par ordinateur -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Computer Aided Design
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    FreeCAD / 
                                                    Skidl
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Conception assistée par ordinateur -->

                            <br>

                            <!-- Gestion de version -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Version management
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Git / 
                                                    GitHub : https://github.com/Jay4C
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Gestion de version -->

                            <br>

                            <!-- Réseaux et Systèmes -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Networks and systems
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Windows / 
                                                    Linux
                                                </li>
                                                <li>
                                                    IT Protocols [HTTP, HTTPS, TCP, IPv4, IPv6, SSH, FTP, 
                                                    SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, Modbus RTU]
                                                </li>
                                                <li>
                                                    MinimalModbus with Python / 
                                                    Jlibmodbus with Java
                                                </li>
                                                <li>
                                                    Routeur 4G / 
                                                    Port Forwarding / 
                                                    DynDNS / 
                                                    Dongle 3G
                                                </li>
                                                <li>
                                                    Wireshark / 
                                                    PuTTY / 
                                                    WinSCP /
                                                    Postman /
                                                    TeamViewer / 
                                                    VirtualBox /
                                                    Tor /
                                                    Angry IP Scanner
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Réseaux et Systèmes -->

                            <br>

                            <!-- Gestion de projet et communications -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Project management and communications
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Brainstorming / 
                                                    FreeMind / 
                                                    Pencil / 
                                                    UML diagrams/ 
                                                    ArgoUML
                                                </li>
                                                <li>
                                                    GanttProject / 
                                                    Planning of a stages of a project
                                                </li>
                                                <li>
                                                    Project management /
                                                    Operations management /
                                                    Agile management /
                                                    Quality management
                                                </li>
                                                <li>
                                                    Customer needs analysis / 
                                                    Establishment of a specifications book
                                                </li>
                                                <li>
                                                    Writing a technical support / 
                                                    Writing a knowledge database
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Gestion de projet et communications -->

                            <br>

                            <!-- Droit -->
                            <table class="table table-bordered">
                                <thead>
                                    <tr>
                                        <th scope="col">
                                            Law
                                        </th>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Tax code in France / 
                                                    Monetary and financial code in France / 
                                                    Environmental code in France / 
                                                    Energy code in France
                                                </li>
                                                <li>
                                                    Code of relations between the public and the administration in France / 
                                                    Code of construction and housing in France
                                                </li>
                                                <li>
                                                    Code of urbanism in France /
                                                    Code of posts and electronic communications in France /                                                 
                                                </li>
                                                <li>
                                                    Website "www.legifrance.gouv.fr" /
                                                    Website "www.lexadin.nl" /
                                                    Website "https://europa.eu/european-union/law/find-legislation_fr"
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Droit -->
                        </div>
                    </div>
                    <!-- Compétences -->

                    <br>

                    <!-- Expériences -->
                    <div class="card">
                        <div class="card-header">
                            <h4>
                                Experiences
                            </h4>
                        </div>
                        <div class="card-body">
                            <!-- Holomorphe -->
                            <table class="table table-striped table-bordered">
                                <thead>
                                    <tr>
                                        <td scope="col">
                                            Title: Versatile IT development
                                            <br>
                                            Period: May 2020 - in progress
                                            <br>
                                            Professional status: President
                                            <br>
                                            Company: Holomorphe - Paris
                                        </td>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Python application development to create mechanical parts
                                                    precision for hydrogen technologies with software
                                                    FreeCad available on https://github.com/jay4c/python-macros-for-freecad
                                                </li>
                                                <li>
                                                    Python application development for robotic process automation
                                                    on https://github.com/Jay4C/Web-Automation
                                                </li>
                                                <li>
                                                    Development of a Python application to recover all suitable plots 
                                                    for injecting synthesis methane into a natural gas distribution 
                                                    network in France
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Holomorphe -->

                            <br>

                            <!-- Société PASS Technologie -->
                            <table class="table table-striped table-bordered">
                                <thead>
                                    <tr>
                                        <td scope="col">
                                            Title: Versatile computer developer
                                            <br>
                                            Period: February 2019 - April 2019 [3 months]
                                            <br>
                                            Professional status: employee
                                            <br>
                                            Company: Pass Technology - Paris
                                        </td>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Development of an application for the BIM (Building Information 
                                                    Modeling)
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Société Pass Tech -->

                            <br>

                            <!-- Kriir -->
                            <table class="table table-striped table-bordered">
                                <thead>
                                    <tr>
                                        <td scope="col">
                                            Title: IT engineer
                                            <br>
                                            Period: September - November 2018 [3 months]
                                            <br>
                                            Professional status: employee
                                            <br>
                                            Company: Kriir - Ermont
                                        </td>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Development of a Java web application to use the electronic signature
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Kriir -->

                            <br>

                            <!-- Refuel -->
                            <table class="table table-striped table-bordered">
                                <thead>
                                    <tr>
                                        <td scope="col">
                                            Title: Reading and writing on an automaton in Modbus
                                            <br>
                                            Period: June - July 2018 [2 months]
                                            <br>
                                            Professional status: Freelance
                                            <br>
                                            Company: Refuel S.A.S. - Trappes
                                        </td>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Development of several features to read and to write on a programmable 
                                                    controller via the Modbus TCP protocol with Java language in WiFi link
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Refuel  -->

                            <br>

                            <!-- SETEC SMART EFFICIENCY -->
                            <table class="table table-striped table-bordered">
                                <thead>
                                    <tr>
                                        <td scope="col">
                                            Title: Automatic generation of energy report
                                            <br>
                                            Period: November - December 2017 [2 months]
                                            <br>
                                            Professional status: Freelance
                                            <br>
                                            Company: SETEC SMART EFFICIENCY - Paris
                                        </td>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Develop a Java program to generate an energy report
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- SETEC SMART EFFICIENCY -->

                            <br>

                            <!--  SOCIÉTÉ GAMELOFT -->
                            <table class="table table-striped table-bordered">
                                <thead>
                                    <tr>
                                        <td scope="col">
                                            Title: Verification of the relevance of data
                                            <br>
                                            Period: June - July 2017 [2 months]
                                            <br>
                                            Professional status: Freelance
                                            <br>
                                            Company: Gameloft - Paris
                                        </td>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Development of a Visual Basic macro for an Excel file to control the 
                                                    consistency of the data entered in several Excel files and consolidate 
                                                    these data in one Excel file for human resources
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- SOCIÉTÉ GAMELOFT -->

                            <br>
                        </div>
                    </div>
                    <!-- Expériences -->

                    <br>

                    <!-- Formations -->
                    <div class="card">
                        <div class="card-header">
                            <h4>
                                Trainings
                            </h4>
                        </div>
                        <div class="card-body">
                            <!-- Formation Java -->
                            <table class="table table-striped table-bordered">
                                <thead>
                                    <tr>
                                        <td scope="col">
                                            Title: Java / JEE training
                                            <br>
                                            Period: April - July 2017 [4 months]
                                            <br>
                                            Training center: INTI Training - Paris Tour Montparnasse
                                        </td>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    Learning of Java IT Language /
                                                    Learning of the object-oriented design
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Formation Java -->

                            <br>

                            <!-- Diplome ingenieur -->
                            <table class="table table-striped table-bordered">
                                <thead>
                                    <tr>
                                        <td scope="col">
                                            Title: General engineer diploma
                                            <br>
                                            Period: September 2013 - August 2016 [3 years]
                                            <br>
                                                Training center: Superior School of Engineers Léonard de Vinci 
                                                - Paris La Défense
                                        </td>
                                    </tr>
                                </thead>
                                <tbody>
                                    <tr>
                                        <td>
                                            <ul>
                                                <li>
                                                    New energies / Computing / Finance / Digital Mechanics
                                                </li>
                                            </ul>
                                        </td>
                                    </tr>
                                </tbody>
                            </table>
                            <!-- Diplome ingenieur -->
                        </div>
                    </div>
                    <!-- Formations -->

                    <br>

                    <!-- Centre d'intérêts -->
                    <div class="card">
                        <div class="card-header">
                            <h4>
                                Center of interests
                            </h4>
                        </div>
                        <div class="card-body">
                            <ul>
                                <li>
                                    Technology /
                                    Patents
                                </li>
                                <li>
                                    French law / 
                                    European law / 
                                    International law
                                </li>
                                <li>
                                    Physics / 
                                    Mathematics / 
                                    Finance / 
                                    Economy / 
                                    Management / 
                                    Marketing / 
                                    Spirituality / 
                                    Ufology
                                </li>
                            </ul>
                        </div>
                    </div>
                    <!-- Centre d'intérêts -->
                </div>
            </div>
        </div>

        <br>

        <!-- Optional JavaScript -->
        <!-- jQuery first, then Popper.js, then Bootstrap JS -->
        <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
            integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
            crossorigin="anonymous"></script>
        <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
            integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
            crossorigin="anonymous"></script>
        <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
            integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
            crossorigin="anonymous"></script>
    </body>

    </html>
            """

        options = {
            'page-size': 'A4',
            'header-center': 'Resume of Mr. ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Holomorphe company [SIREN : 883632556]'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body, "CV\\Resume_Of_Mr_.pdf", configuration=config,
                           options=options)

    # ok
    def test_resume_renewable_energy_engineer(self):
        print("test_resume_renewable_energy_engineer")

        body = """
        <!doctype html>
        <html lang="en">
    
        <head>
            <!-- Required meta tags -->
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">
    
            <!-- Bootstrap CSS -->
            <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/css/bootstrap.min.css"
                integrity="sha384-MCw98/SFnGE8fJT3GXwEOngsV7Zt27NXFoaoApmYm81iuXoPkFOJwJ8ERdknLPMO" crossorigin="anonymous">
    
            <title>Resume of  - Renewable energy engineer</title>
        </head>
    
        <body>
            <br>
            <div class="container">
                <div class="card text-center">
                    <div class="card-header text-center">
                        <h2>
                            Renewable energy engineer
                        </h2>
                    </div>
                    <div class="card-body text-center">
    
                        <!-- Informations personnelles -->
                        <table class="table table-striped table-bordered">
                            <thead>
                                <tr>
                                    <th scope="col">Identity</th>
                                    <th scope="col">Postal address</th>
                                    <th scope="col">Phone</th>
                                    <th scope="col">Email</th>
                                    <th scope="col">Nationality</th>
                                </tr>
                            </thead>
                            <tbody>
                                <tr>
                                    <td>
                                        
                                    </td>
                                    <td>
                                         93800 Epinay-sur-Seine (France)
                                    </td>
                                    <td>
                                        
                                    </td>
                                    <td>
                                        @outlook.fr
                                    </td>
                                    <td>
                                        French
                                    </td>
                                </tr>
                            </tbody>
                        </table>
                        <!-- Informations personnelles -->
    
                        <br>
    
                        <!--  Compétences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Skills
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Energy -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Energy
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- Type of machines -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Type of machines
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Photovoltaic panels / Wind turbine / Gas turbine / 
                                                                Nuclear fusion / Cold fusion / Free energy devices
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Type of machines -->
                                        
                                        <br>
                                        
                                        <!-- Mathematics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mathematics
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Complex numbers / Elementary geometry of the plane / 
                                                                Elementary geometry of space / Usual functions / 
                                                                Linear differential equations / Study of plane curves / 
                                                                Conical / Natural integers, finite sets, counts / 
                                                                Field of real numbers / Sequences of real numbers / 
                                                                Functions of a real-valued real variable / Derivation 
                                                                of real-value functions / Integration of real-value 
                                                                functions on a segment / Limited developments / 
                                                                Metric properties of arcs / Suites and functions with 
                                                                complex values / Notions about the functions of two 
                                                                real variables / Multiple integrals / Algebraic 
                                                                structures / Arithmetic / Polynomials / Rational 
                                                                fractions / Vector spaces / Dimension of vector 
                                                                spaces / Matrix calculation / Symmetric group, 
                                                                determinant / Scalar product, orthogonal group / 
                                                                Affine geometry / Demonstration techniques / Algebra 
                                                                techniques / Analysis techniques / Multi-variable 
                                                                analysis / Graph theory / Information theory / Data 
                                                                analysis
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mathematics -->
                                        
                                        <br>
                                        
                                        <!-- Physics -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Physics
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Thermodynamics / Electrotechnical / Electromagnetism / 
                                                                Electronic / Thermo-electricity / Quantum mechanics / 
                                                                Nuclear physics / Electricity / Plasma physics / 
                                                                Nuclear magnetic resonance / Dielectric materials / 
                                                                Fluid mechanics / Harmonic oscillators and sinusoidal 
                                                                signals / Linear circuits in continuous mode / 
                                                                Transitional arrangements of the first order / 
                                                                Transitional regime of the second order / 
                                                                Forced sinusoidal regime - resonance / Filtering / 
                                                                Kinematics / Fundamental principle of dynamics / 
                                                                Mechanical energy / Motion in an electric and magnetic 
                                                                field / Signal propagation - notion of waves / 
                                                                Snell-Descartes laws - reflection and refraction / 
                                                                Spherical thin lenses / Kinetic moment - solid rotating 
                                                                around a fixed axis / Conservative central forces / 
                                                                States of matter / Thermal machines / Magnetic field 
                                                                - laplace forces - induction
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Physics -->
                                        
                                        <br>
                                        
                                        <!-- Management -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Management
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Collaborative management / Sales management / Advisory 
                                                                management / Team management / Innovation management / 
                                                                Project Management / Quality management / Knowledge 
                                                                management / Operations management / Human resources 
                                                                management / Risk management / Change management / 
                                                                Event management
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Management -->
                                        
                                        <br>
                                        
                                        <!-- Marketing -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Marketing
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Commercial marketing / B2B marketing / Digital 
                                                                marketing / Direct marketing / Operational marketing / 
                                                                Relationship marketing / Strategic marketing / 
                                                                Event marketing
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Marketing -->
                                        
                                        <br>
                                        
                                        <!-- Finance -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Finance
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Financial analysis / Corporate finance / International 
                                                                finance / Personal finance / Project finance / Forex
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Finance -->
                                        
                                        <br>
                                        
                                        <!-- Economy -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Economy
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Macroeconomics / Microeconomics / Economic agents / 
                                                                Agricultural economics / Circular economy / Consumption 
                                                                and savings / Economics of education / Financial 
                                                                economy / Geographical economy / Industrial economics 
                                                                / Global economy / Monetary economy
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Economy -->
                                        
                                        <br>
                                        
                                        <!-- International law -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        International law
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Website : "www.legifrance.gouv.fr"
                                                            </li>
                                                            <li>
                                                                Website : "www.lexadin.nl"
                                                            </li>
                                                            <li>
                                                                Website : "https://europa.eu/european-union/law/find-legislation_fr"
                                                            </li>
                                                            <li>
                                                                Website : "https://www.fedlex.admin.ch/fr/home"
                                                            </li>
                                                            <li>
                                                                Website : "https://legislation.gov.ky/"
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- International law -->        
                                    </div>
                                </div>
                                <!-- Energy -->
                                
                                <br>
                                
                                <!-- Computing -->
                                <div class="card">
                                    <div class="card-header">
                                        <h4>
                                            Computing
                                        </h4>
                                    </div>
                                    <div class="card-body">
                                        <!-- preferred programming languages -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Preferred programming languages
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>Python / Java / JavaScript</li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Preferred programming languages -->    
                                        
                                        <br>
                                        
                                        <!-- Office -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Office
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Microsoft Excel, Word, PowerPoint, Access, 
                                                                Outlook
                                                            </li>
                                                            <li>
                                                                Open Office Calc, Writer
                                                            </li>
                                                            <li>
                                                                Intellij for Java, PyCharm for Python
                                                            </li>
                                                            <li>
                                                                Pencil software to make simplified schemes / 
                                                                GanttProject software for project management
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Office -->
                                        
                                        <br>
                                        
                                        <!-- Front-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Front-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                HTML markup language /
                                                                CSS /
                                                                JQuery Framework /
                                                                Django Framework /
                                                                MJML Framework for Email Responsive
                                                            </li>
                                                            <li>
                                                                Visual Studio Software Code /
                                                                Angular Framework 7 /
                                                                Boostrap /
                                                                Leaflet.js
                                                            </li>
                                                            <li>
                                                                Vis.js / 
                                                                Artyom.js /
                                                                Flowchart.js
                                                            </li>
                                                            <li>
                                                                Tomcat server /
                                                                Apache server /
                                                                Nginx server
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Front-end -->
            
                                        <br>
            
                                        <!-- Back-end -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Back-end
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Python /
                                                                Flask /
                                                                BeautifulSoup /
                                                                pdfkit /
                                                                xlswriter /
                                                                xlrd /
                                                                pywinauto /
                                                                Hyperledger Sawtooth
                                                            </li>
                                                            <li>
                                                                Java /
                                                                Maven / 
                                                                Gradle /
                                                                Spring /
                                                                Jsoup /
                                                                Selenium WebDriver
                                                            </li>
                                                            <li>
                                                                JavaScript /
                                                                Visual Basic /
                                                                C# /
                                                                TypeScript
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Back-end -->
            
                                        <br>
            
                                        <!-- Tests unitaires et qualité -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Unit and quality tests
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Unittests with Python / 
                                                                JUnit with Java / 
                                                                Mockito with Java / 
                                                                WireMock with Java
                                                            </li>
                                                            <li>
                                                                SonarQube
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Tests unitaires et qualité -->
            
                                        <br>
            
                                        <!-- Systèmes de bases de données -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Database systems
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                MySQL / 
                                                                MySQL Workbench
                                                            </li>
                                                            <li>
                                                                NeO4J, graph database server  / 
                                                                MongoDB, documents database server
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes de bases de données -->
            
                                        <br>
            
                                        <!-- Outils de déploiement -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        DevOps
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Docker / Git
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Outils de déploiement -->
            
                                        <br>
            
                                        <!-- Intelligence artificielle -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Artificial intelligence
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Data mining / 
                                                                Knowledge management /
                                                                Robotic Process Automation
                                                            </li>
                                                            <li>
                                                                Cloud computing [Open data; Autodesk Forge; Netheos] / 
                                                                Speech recognition / 
                                                                Computer vision
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Intelligence artificielle -->
            
                                        <br>
            
                                        <!-- Mobile -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Mobile
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Android /
                                                                Kivy
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Mobile -->
            
                                        <br>
            
                                        <!-- Systèmes embarqués -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Embedded systems
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Raspberry Pi
                                                            </li>
                                                            <li>
                                                                Arduino Uno and Arduino Yun /
                                                                OpenWRT OS
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Systèmes embarqués -->
            
                                        <br>
            
                                        <!-- Conception assistée par ordinateur -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Computer Aided Design
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                FreeCAD / 
                                                                Skidl
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Conception assistée par ordinateur -->
            
                                        <br>
            
                                        <!-- Gestion de version -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Version management
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Git / 
                                                                GitHub : https://github.com/Jay4C
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de version -->
            
                                        <br>
            
                                        <!-- Réseaux et Systèmes -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Networks and systems
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Windows / 
                                                                Linux
                                                            </li>
                                                            <li>
                                                                IT Protocols [HTTP, HTTPS, TCP, IPv4, IPv6, SSH, FTP, 
                                                                SFTP, SCP, SMTP, IMAP, DHCP, Modbus TCP, Modbus RTU]
                                                            </li>
                                                            <li>
                                                                MinimalModbus with Python / 
                                                                Jlibmodbus with Java
                                                            </li>
                                                            <li>
                                                                Routeur 4G / 
                                                                Port Forwarding / 
                                                                DynDNS / 
                                                                Dongle 3G
                                                            </li>
                                                            <li>
                                                                Wireshark / 
                                                                PuTTY / 
                                                                WinSCP /
                                                                Postman /
                                                                TeamViewer / 
                                                                VirtualBox /
                                                                Tor /
                                                                Angry IP Scanner
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Réseaux et Systèmes -->
            
                                        <br>
            
                                        <!-- Gestion de projet et communications -->
                                        <table class="table table-bordered">
                                            <thead>
                                                <tr>
                                                    <th scope="col">
                                                        Project management and communications
                                                    </th>
                                                </tr>
                                            </thead>
                                            <tbody>
                                                <tr>
                                                    <td>
                                                        <ul>
                                                            <li>
                                                                Brainstorming / 
                                                                FreeMind / 
                                                                Pencil / 
                                                                UML diagrams/ 
                                                                ArgoUML
                                                            </li>
                                                            <li>
                                                                GanttProject / 
                                                                Planning of a stages of a project
                                                            </li>
                                                            <li>
                                                                Project management /
                                                                Operations management /
                                                                Agile management /
                                                                Quality management
                                                            </li>
                                                            <li>
                                                                Customer needs analysis / 
                                                                Establishment of a specifications book
                                                            </li>
                                                            <li>
                                                                Writing a technical support / 
                                                                Writing a knowledge database
                                                            </li>
                                                        </ul>
                                                    </td>
                                                </tr>
                                            </tbody>
                                        </table>
                                        <!-- Gestion de projet et communications -->        
                                    </div>
                                </div>
                                <!-- Computing -->
                            </div>
                        </div>
                        <!-- Compétences -->
    
                        <br>
    
                        <!-- Expériences -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Experiences
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Société PASS Technologie -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Title: Versatile computer developer
                                                <br>
                                                Period: February 2019 - April 2019 [3 months]
                                                <br>
                                                Professional status: employee
                                                <br>
                                                Company: Pass Technology - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Development of an application for the BIM (Building Information 
                                                        Modeling)
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Société Pass Tech -->
    
                                <br>
    
                                <!-- Kriir -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Title: IT engineer
                                                <br>
                                                Period: September - November 2018 [3 months]
                                                <br>
                                                Professional status: employee
                                                <br>
                                                Company: Kriir - Ermont
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Development of a Java web application to use the electronic signature
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Kriir -->
    
                                <br>
    
                                <!-- Refuel -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Title: Reading and writing on an automaton in Modbus
                                                <br>
                                                Period: June - July 2018 [2 months]
                                                <br>
                                                Professional status: Freelance
                                                <br>
                                                Company: Refuel S.A.S. - Trappes
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Development of several features to read and to write on a programmable 
                                                        controller via the Modbus TCP protocol with Java language in WiFi link
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Refuel  -->
    
                                <br>
    
                                <!-- SETEC SMART EFFICIENCY -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Title: Automatic generation of energy report
                                                <br>
                                                Period: November - December 2017 [2 months]
                                                <br>
                                                Professional status: Freelance
                                                <br>
                                                Company: SETEC SMART EFFICIENCY - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Develop a Java program to generate an energy report
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SETEC SMART EFFICIENCY -->
    
                                <br>
    
                                <!--  SOCIÉTÉ GAMELOFT -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Title: Verification of the relevance of data
                                                <br>
                                                Period: June - July 2017 [2 months]
                                                <br>
                                                Professional status: Freelance
                                                <br>
                                                Company: Gameloft - Paris
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Development of a Visual Basic macro for an Excel file to control the 
                                                        consistency of the data entered in several Excel files and consolidate 
                                                        these data in one Excel file for human resources
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- SOCIÉTÉ GAMELOFT -->
                            </div>
                        </div>
                        <!-- Expériences -->
    
                        <br>
    
                        <!-- Formations -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Trainings
                                </h4>
                            </div>
                            <div class="card-body">
                                <!-- Formation Java -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Title: Java / JEE training
                                                <br>
                                                Period: April - July 2017 [4 months]
                                                <br>
                                                Training center: INTI Training - Paris Tour Montparnasse
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        Learning of Java IT Language /
                                                        Learning of the object-oriented design
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Formation Java -->
    
                                <br>
    
                                <!-- Diplome ingenieur -->
                                <table class="table table-striped table-bordered">
                                    <thead>
                                        <tr>
                                            <td scope="col">
                                                Title: General engineer diploma
                                                <br>
                                                Period: September 2013 - August 2016 [3 years]
                                                <br>
                                                    Training center: Superior School of Engineers Léonard de Vinci 
                                                    - Paris La Défense in France
                                            </td>
                                        </tr>
                                    </thead>
                                    <tbody>
                                        <tr>
                                            <td>
                                                <ul>
                                                    <li>
                                                        New energies / Computing / Finance / Digital Mechanics
                                                    </li>
                                                </ul>
                                            </td>
                                        </tr>
                                    </tbody>
                                </table>
                                <!-- Diplome ingenieur -->
                            </div>
                        </div>
                        <!-- Formations -->
    
                        <br>
                        
                        <!-- Languages -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Languages
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        English written and oral [TOEIC: 840 points; Expiry date: July 11, 2020]
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Languages -->
                        
                        <br>
    
                        <!-- Centre d'intérêts -->
                        <div class="card">
                            <div class="card-header">
                                <h4>
                                    Center of interests
                                </h4>
                            </div>
                            <div class="card-body">
                                <ul>
                                    <li>
                                        Technology /
                                        Patents
                                    </li>
                                    <li>
                                        French law / 
                                        European law / 
                                        International law
                                    </li>
                                    <li>
                                        Physics / 
                                        Mathematics / 
                                        Finance / 
                                        Economy / 
                                        Management / 
                                        Marketing / 
                                        Spirituality / 
                                        Ufology
                                    </li>
                                </ul>
                            </div>
                        </div>
                        <!-- Centre d'intérêts -->
                    </div>
                </div>
            </div>
    
            <br>
    
            <!-- Optional JavaScript -->
            <!-- jQuery first, then Popper.js, then Bootstrap JS -->
            <script src="https://code.jquery.com/jquery-3.3.1.slim.min.js"
                integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo"
                crossorigin="anonymous"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/popper.js/1.14.3/umd/popper.min.js"
                integrity="sha384-ZMP7rVo3mIykV+2+9J3UJ46jBk0WLaUAdn689aCwoqbBJiSnjAK/l8WvCWPIPm49"
                crossorigin="anonymous"></script>
            <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.1.3/js/bootstrap.min.js"
                integrity="sha384-ChfqqxuZUCnJSK3+MXmPNIyE6ZbWh2IMqE241rYiqJxyMiZ6OW/JmZQ5stwEULTy"
                crossorigin="anonymous"></script>
        </body>
        </html>
        """

        options = {
            'page-size': 'A4',
            'header-center': 'Resume of ',
            'footer-right': '[page] sur [topage]',
            'footer-left': 'Renewable energy engineer'
        }

        path_wkthmltopdf = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
        config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
        pdfkit.from_string(body, "CV\\Resume_Of__[Renewable_energy_engineer].pdf", configuration=config,
                           options=options)


if __name__ == '__main__':
    unittest.main()
